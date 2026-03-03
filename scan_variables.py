#!/usr/bin/env python3
"""
扫描立案文件夹中的所有文档，提取 {变量}
"""
import os
import re
from docx import Document

def scan_docx_variables(filepath):
    """扫描 Word 文档中的 {变量}"""
    variables = set()
    try:
        doc = Document(filepath)
        # 扫描段落
        for para in doc.paragraphs:
            matches = re.findall(r'\{([^}]+)\}', para.text)
            variables.update(matches)
        # 扫描表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = re.findall(r'\{([^}]+)\}', cell.text)
                    variables.update(matches)
    except Exception as e:
        print(f"  读取失败: {e}")
    return variables

def scan_directory(base_path):
    """递归扫描目录"""
    all_variables = {}
    
    for root, dirs, files in os.walk(base_path):
        for file in files:
            if file.startswith('~') or file.startswith('.'):
                continue
            if file.endswith('.docx'):
                filepath = os.path.join(root, file)
                rel_path = os.path.relpath(filepath, base_path)
                print(f"\n扫描: {rel_path}")
                vars = scan_docx_variables(filepath)
                if vars:
                    print(f"  发现变量: {vars}")
                    all_variables[rel_path] = vars
    
    return all_variables

if __name__ == '__main__':
    base_path = '文件生成/1-立案'
    print(f"扫描目录: {base_path}")
    print("=" * 60)
    
    result = scan_directory(base_path)
    
    print("\n" + "=" * 60)
    print("汇总所有变量:")
    all_vars = set()
    for file_vars in result.values():
        all_vars.update(file_vars)
    
    for var in sorted(all_vars):
        print(f"  - {{{var}}}")
