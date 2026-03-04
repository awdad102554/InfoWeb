#!/usr/bin/env python3
"""
批量文档生成器 - 支持多文件生成和合并
复用 document_generator.py 的生成逻辑，确保格式完全一致
"""
import os
import re
import io
import tempfile
import zipfile
from datetime import datetime
from copy import deepcopy

# 导入原来的文档生成器
from document_generator import DocumentGenerator


class BatchDocumentGenerator:
    """批量文档生成器"""
    
    def __init__(self, doc_templates_dir):
        self.doc_templates_dir = doc_templates_dir
    
    def generate_batch(self, template_paths, case_data, case_no):
        """
        批量生成文档并合并
        
        Args:
            template_paths: 模板路径列表
            case_data: 案件数据（原始API返回数据）
            case_no: 案件编号（用于文件名）
        
        Returns:
            dict: {
                'docx': {'path': str, 'files': list},
                'xlsx': {'path': str, 'files': list},
            }
        """
        # 按类型分组
        docx_files = []
        xlsx_files = []
        
        for path in template_paths:
            ext = os.path.splitext(path)[1].lower()
            if ext == '.docx':
                docx_files.append(path)
            elif ext in ['.xls', '.xlsx']:
                xlsx_files.append(path)
        
        result = {}
        
        # 生成并合并docx
        if docx_files:
            docx_path = self._generate_and_merge_docx(docx_files, case_data, case_no)
            result['docx'] = {'path': docx_path, 'files': docx_files}
        
        # 生成并合并xlsx
        if xlsx_files:
            xlsx_path = self._generate_and_merge_xlsx(xlsx_files, case_data, case_no)
            result['xlsx'] = {'path': xlsx_path, 'files': xlsx_files}
        
        return result
    
    def _generate_and_merge_docx(self, template_paths, case_data, case_no):
        """生成并合并多个docx文件"""
        from docx import Document
        
        # 创建临时目录
        temp_dir = tempfile.mkdtemp()
        
        # 使用原来的 DocumentGenerator 生成所有文档到临时文件
        generated_files = []
        for i, template_path in enumerate(template_paths):
            full_path = os.path.join(self.doc_templates_dir, template_path)
            if not os.path.exists(full_path):
                continue
            
            temp_output = os.path.join(temp_dir, f'temp_{i}.docx')
            
            # 使用原来的 DocumentGenerator
            generator = DocumentGenerator(full_path, temp_output)
            generator.generate(case_data)
            
            generated_files.append(temp_output)
        
        if not generated_files:
            raise ValueError("没有成功生成任何docx文件")
        
        # 合并所有docx
        output_filename = self._generate_output_filename(case_no, '.docx')
        output_path = os.path.join(self.doc_templates_dir, 'output', output_filename)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # 使用python-docx合并，保留格式
        self._merge_docx_files(generated_files, output_path)
        
        # 清理临时文件
        for f in generated_files:
            try:
                os.remove(f)
            except:
                pass
        try:
            os.rmdir(temp_dir)
        except:
            pass
        
        return output_path
    
    def _generate_and_merge_xlsx(self, template_paths, case_data, case_no):
        """生成并合并多个xlsx文件（每个文件作为一个sheet）"""
        # 创建临时目录
        temp_dir = tempfile.mkdtemp()
        
        # 使用原来的 DocumentGenerator 生成所有文档到临时文件
        generated_files = []
        for i, template_path in enumerate(template_paths):
            full_path = os.path.join(self.doc_templates_dir, template_path)
            if not os.path.exists(full_path):
                continue
            
            temp_output = os.path.join(temp_dir, f'temp_{i}.xlsx')
            
            # 使用原来的 DocumentGenerator
            generator = DocumentGenerator(full_path, temp_output)
            generator.generate(case_data)
            
            generated_files.append((temp_output, template_path))
        
        if not generated_files:
            raise ValueError("没有成功生成任何xlsx文件")
        
        # 合并所有xlsx
        output_filename = self._generate_output_filename(case_no, '.xlsx')
        output_path = os.path.join(self.doc_templates_dir, 'output', output_filename)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # 如果是单个文件，直接复制
        if len(generated_files) == 1:
            import shutil
            shutil.copy(generated_files[0][0], output_path)
        else:
            # 合并多个xlsx
            self._merge_xlsx_files(generated_files, output_path)
        
        # 清理临时文件
        for f, _ in generated_files:
            try:
                os.remove(f)
            except:
                pass
        try:
            os.rmdir(temp_dir)
        except:
            pass
        
        return output_path
    
    def _merge_docx_files(self, file_paths, output_path):
        """合并多个docx文件
        使用 docxcompose 库进行专业合并，每个文档之间添加分页符
        如果 docxcompose 失败，使用备用方案手动复制内容
        """
        from docx import Document
        from docx.enum.text import WD_BREAK
        
        # 首先尝试使用 docxcompose（保留格式更好）
        try:
            from docxcompose.composer import Composer
            
            # 使用第一个文档作为基础
            master = Document(file_paths[0])
            composer = Composer(master)
            
            # 添加后续文档，每个文档前添加分页符
            for file_path in file_paths[1:]:
                master.add_page_break()
                doc = Document(file_path)
                composer.append(doc)
            
            composer.save(output_path)
            return
        except Exception as e:
            # docxcompose 失败，使用备用方案
            print(f"docxcompose 合并失败，使用备用方案: {e}")
        
        # 备用方案：手动复制内容
        self._merge_docx_files_fallback(file_paths, output_path)
    
    def _merge_docx_files_fallback(self, file_paths, output_path):
        """备用合并方案 - 手动复制内容
        当 docxcompose 失败时使用，兼容性更好但格式保留可能不完整
        """
        from docx import Document
        from docx.enum.text import WD_BREAK
        from copy import deepcopy
        
        # 使用第一个文档作为基础
        master = Document(file_paths[0])
        
        # 添加后续文档
        for file_path in file_paths[1:]:
            # 添加分页符
            master.add_page_break()
            
            # 读取文档
            doc = Document(file_path)
            
            # 复制所有段落
            for para in doc.paragraphs:
                new_para = master.add_paragraph()
                # 复制段落文本
                new_para.text = para.text
                # 尝试复制段落样式
                try:
                    new_para.style = para.style
                except:
                    pass
                # 复制段落格式
                try:
                    new_para.alignment = para.alignment
                except:
                    pass
            
            # 复制所有表格
            for table in doc.tables:
                # 创建新表格
                new_table = master.add_table(
                    rows=len(table.rows),
                    cols=len(table.columns)
                )
                
                # 复制单元格内容
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                            new_cell = new_table.rows[i].cells[j]
                            # 复制文本
                            new_cell.text = cell.text
                            # 尝试复制单元格格式
                            try:
                                if cell.vertical_alignment:
                                    new_cell.vertical_alignment = cell.vertical_alignment
                            except:
                                pass
        
        master.save(output_path)
    
    def _merge_xlsx_files(self, file_paths_with_names, output_path):
        """合并多个xlsx文件，每个文件作为一个sheet
        每个文件只保留第一个工作表，重命名后合并到新工作簿
        尽可能保留原格式（字体、边框、对齐等）
        自动跳过开头的空行
        """
        from openpyxl import Workbook, load_workbook
        from copy import copy
        
        # 创建新的空工作簿
        merged_wb = Workbook()
        merged_wb.remove(merged_wb.active)  # 删除默认sheet
        
        # 逐个添加文件（每个文件作为一个sheet）
        for idx, (file_path, original_name) in enumerate(file_paths_with_names, 1):
            try:
                wb = load_workbook(file_path, data_only=False, keep_vba=False)
                
                # 只复制第一个工作表
                source_sheet = wb[wb.sheetnames[0]]
                
                # 创建新sheet名称
                base_name = os.path.splitext(os.path.basename(original_name))[0]
                new_sheet_name = f"{idx}_{base_name}"[:31]
                
                # 创建工作表
                target_sheet = merged_wb.create_sheet(title=new_sheet_name)
                
                # 找到第一个有内容的行（跳过开头的空行）
                first_data_row = 1
                for row_idx in range(1, source_sheet.max_row + 1):
                    has_content = False
                    for col_idx in range(1, source_sheet.max_column + 1):
                        cell = source_sheet.cell(row=row_idx, column=col_idx)
                        if cell.value is not None and str(cell.value).strip():
                            has_content = True
                            break
                    if has_content:
                        first_data_row = row_idx
                        break
                
                # 计算行偏移量（用于调整行高复制）
                row_offset = first_data_row - 1
                
                # 从第一个有内容的行开始复制
                target_row_idx = 1
                for row_idx in range(first_data_row, source_sheet.max_row + 1):
                    for col_idx in range(1, source_sheet.max_column + 1):
                        source_cell = source_sheet.cell(row=row_idx, column=col_idx)
                        new_cell = target_sheet.cell(row=target_row_idx, column=col_idx)
                        
                        # 复制值
                        if source_cell.value is not None:
                            new_cell.value = source_cell.value
                        
                        # 复制格式（如果单元格有样式）
                        if source_cell.has_style:
                            # 复制数字格式
                            if source_cell.number_format:
                                new_cell.number_format = source_cell.number_format
                            
                            # 复制字体
                            if source_cell.font:
                                new_cell.font = copy(source_cell.font)
                            
                            # 复制边框
                            if source_cell.border:
                                new_cell.border = copy(source_cell.border)
                            
                            # 复制填充/背景
                            if source_cell.fill:
                                new_cell.fill = copy(source_cell.fill)
                            
                            # 复制对齐
                            if source_cell.alignment:
                                new_cell.alignment = copy(source_cell.alignment)
                            
                            # 复制保护
                            if source_cell.protection:
                                new_cell.protection = copy(source_cell.protection)
                    
                    # 复制行高（调整行号）
                    if row_idx in source_sheet.row_dimensions:
                        row_dim = source_sheet.row_dimensions[row_idx]
                        if row_dim.height:
                            target_sheet.row_dimensions[target_row_idx].height = row_dim.height
                    
                    target_row_idx += 1
                
                # 复制列宽
                for col_letter, col_dim in source_sheet.column_dimensions.items():
                    if col_dim.width:
                        target_sheet.column_dimensions[col_letter].width = col_dim.width
                
                # 复制页面设置（如打印区域、标题行等）
                if source_sheet.print_title_rows:
                    target_sheet.print_title_rows = source_sheet.print_title_rows
                if source_sheet.print_area:
                    target_sheet.print_area = source_sheet.print_area
                
                # 复制sheet的显示选项
                target_sheet.sheet_format = copy(source_sheet.sheet_format)
                target_sheet.sheet_properties = copy(source_sheet.sheet_properties)
                
                # 复制合并单元格（调整行号）
                from openpyxl.utils import get_column_letter
                for merged_range in source_sheet.merged_cells.ranges:
                    # 调整合并区域的行号
                    if merged_range.min_row >= first_data_row:
                        new_min_row = merged_range.min_row - row_offset
                        new_max_row = merged_range.max_row - row_offset
                        # 创建新的合并区域字符串
                        min_col_letter = get_column_letter(merged_range.min_col)
                        max_col_letter = get_column_letter(merged_range.max_col)
                        new_range = f"{min_col_letter}{new_min_row}:{max_col_letter}{new_max_row}"
                        target_sheet.merge_cells(new_range)
                
                wb.close()
                
            except Exception as e:
                print(f"合并xlsx失败 {original_name}: {e}")
                import traceback
                traceback.print_exc()
                continue
        
        if len(merged_wb.sheetnames) == 0:
            raise ValueError("没有成功合并任何xlsx文件")
        
        merged_wb.save(output_path)
    
    def _generate_output_filename(self, case_no, ext):
        """生成输出文件名"""
        # 清理案号
        if case_no and str(case_no).startswith('明'):
            case_no = str(case_no)[1:]
        safe_case_no = re.sub(r'[\\/:*?"<>|]', '_', str(case_no))
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        return f"{safe_case_no}-{timestamp}{ext}"


if __name__ == '__main__':
    # 测试代码
    import sys
    if len(sys.argv) > 1:
        generator = BatchDocumentGenerator('文件生成')
        # 测试数据（模拟API返回格式）
        test_data = {
            'data': {
                'cases': {'case_no': '测试案号', 'case_reason': '测试案由'},
                'applicant_arr': [{'name': '测试申请人'}],
                'respondent_arr': [{'name': '测试被申请人'}]
            }
        }
        result = generator.generate_batch(sys.argv[1:], test_data, '测试案号')
        print(f"生成结果: {result}")
