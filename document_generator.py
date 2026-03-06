#!/usr/bin/env python3
"""
仲裁文书生成器 - 根据案件数据填充模板
支持 Word (.docx) 和 Excel (.xls, .xlsx) 格式
"""
import os
import re
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocumentGenerator:
    """文书生成器"""
    
    def __init__(self, template_path, output_path):
        self.template_path = template_path
        self.output_path = output_path
        self.file_ext = os.path.splitext(template_path)[1].lower()
        
    def generate(self, data, target_applicant=None, target_applicants=None):
        """
        根据数据填充模板
        data: 立案详情API返回的数据
        target_applicant: 目标申请人姓名（可选，如果指定则只生成该申请人的文书）- 单个申请人，已废弃，请使用 target_applicants
        target_applicants: 目标申请人姓名列表（可选，如果指定则只使用这些申请人的信息填充）- 支持多申请人
        """
        # 兼容旧版本：如果传了 target_applicant 但没传 target_applicants
        if target_applicant and not target_applicants:
            target_applicants = [target_applicant]
        
        # 预处理数据 - 将复杂表达式转换为简单变量
        processed_data = self._preprocess_data(data, target_applicants=target_applicants)
        
        print(f"预处理后的数据: {processed_data}")
        
        # 根据文件类型选择不同的处理方式
        if self.file_ext == '.docx':
            self._generate_word(processed_data)
        elif self.file_ext in ['.xls', '.xlsx']:
            self._generate_excel(processed_data)
        else:
            raise ValueError(f"不支持的文件格式: {self.file_ext}")
        
        return self.output_path
    
    def _generate_word(self, data):
        """生成 Word 文档"""
        doc = Document(self.template_path)
        
        # 替换所有段落中的变量
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, data)
        
        # 替换表格中的变量
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, data)
        
        # 替换文本框（textbox）中的变量
        self._replace_in_textboxes(doc, data)
        
        doc.save(self.output_path)
    
    def _replace_in_textboxes(self, doc, data):
        """替换文本框中的变量"""
        from docx.oxml.ns import qn
        
        # 遍历文档中的所有元素，查找文本框内容
        body = doc._element
        
        # 查找所有 txbxContent 元素（文本框内容）
        for txbx in body.iter(qn('w:txbxContent')):
            for para in txbx.findall(qn('w:p')):
                self._replace_in_paragraph_element(para, data)
    
    def _replace_in_paragraph_element(self, p_element, data):
        """替换段落元素中的变量（XML 级别）"""
        from docx.oxml.ns import qn
        
        # 获取段落的所有文本
        text = ''.join(t.text or '' for t in p_element.iter(qn('w:t')))
        if not text:
            return
        
        original_text = text
        new_text = self._replace_text(text, data)
        
        if new_text != original_text:
            # 找到第一个 w:t 元素并替换其文本
            first_t = None
            for t in p_element.iter(qn('w:t')):
                first_t = t
                break
            
            if first_t is not None:
                # 检查是否有下划线或粗体格式需要保留
                has_underline = False
                has_bold = False
                
                for r_elem in p_element.iter(qn('w:r')):
                    # 检查这个 run 是否包含变量文本
                    run_text = ''.join(t.text or '' for t in r_elem.iter(qn('w:t')))
                    if '{' in run_text:
                        # 检查是否有下划线格式
                        if list(r_elem.iter(qn('w:u'))):
                            has_underline = True
                        # 检查是否有粗体格式
                        if list(r_elem.iter(qn('w:b'))):
                            has_bold = True
                
                # 清除其他所有 w:t 元素
                parent_map = {}
                for parent in p_element.iter():
                    for child in parent:
                        parent_map[child] = parent
                
                t_elements = list(p_element.iter(qn('w:t')))
                for i, t_elem in enumerate(t_elements):
                    if i == 0:
                        # 第一个元素设置新文本
                        t_elem.text = new_text
                        # 清除 xml:space 属性（如果有）
                        if '{http://www.w3.org/XML/1998/namespace}space' in t_elem.attrib:
                            del t_elem.attrib['{http://www.w3.org/XML/1998/namespace}space']
                    else:
                        # 其他元素清空
                        t_elem.text = ''
                
                # 找到包含第一个 w:t 的 w:r 元素
                first_r = None
                for r_elem in p_element.iter(qn('w:r')):
                    for t_elem in r_elem.iter(qn('w:t')):
                        if t_elem == first_t:
                            first_r = r_elem
                            break
                    if first_r is not None:
                        break
                
                if first_r is not None:
                    # 如果替换了 case_no 相关变量，去除粗体格式
                    if '{case_no' in original_text or '{case_no_raw' in original_text:
                        # 移除所有 w:b 元素（粗体标记）
                        for b_elem in list(first_r.iter(qn('w:b'))):
                            first_r.remove(b_elem)
                    elif has_bold:
                        # 保留粗体格式（添加 w:b 元素）
                        if not list(first_r.iter(qn('w:b'))):
                            from docx.oxml import OxmlElement
                            b_elem = OxmlElement('w:b')
                            first_r.insert(0, b_elem)
                    
                    # 保留下划线格式
                    if has_underline and '{case_no' not in original_text and '{case_no_raw' not in original_text:
                        # 添加下划线元素（如果不存在）
                        if not list(first_r.iter(qn('w:u'))):
                            from docx.oxml import OxmlElement
                            u_elem = OxmlElement('w:u')
                            u_elem.set(qn('w:val'), 'single')  # 单下划线
                            first_r.insert(0, u_elem)
    
    def _generate_excel(self, data):
        """生成 Excel 文档 - 保留格式"""
        if self.file_ext == '.xlsx':
            self._generate_excel_xlsx(data)
        elif self.file_ext == '.xls':
            self._generate_excel_xls(data)
    
    def _generate_excel_xlsx(self, data):
        """生成 .xlsx 格式，严格保留格式"""
        from openpyxl import load_workbook
        
        # 加载工作簿，保留所有格式
        # data_only=False: 保留公式而不是计算值
        # keep_vba=False: 不保留VBA（一般模板没有VBA代码）
        wb = load_workbook(self.template_path, data_only=False, keep_vba=False)
        
        # 遍历所有工作表
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            
            # 遍历所有单元格
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        original_value = cell.value
                        new_value = self._replace_text(original_value, data)
                        if new_value != original_value:
                            # 关键：只修改value，保留所有其他格式属性
                            cell.value = new_value
        
        # 保存时使用相同的文件格式
        wb.save(self.output_path)
    
    def _generate_excel_xls(self, data):
        """生成 .xls 格式，严格保留格式 - 使用低级API完整复制格式"""
        import xlrd
        import xlwt
        from xlutils.copy import copy
        
        # 打开原工作簿，保留格式信息
        rb = xlrd.open_workbook(self.template_path, formatting_info=True)
        
        # 创建一个新的工作簿，手动复制所有格式
        wb = xlwt.Workbook(style_compression=2)
        
        # 复制样式表
        def get_xf_style(xf_index, rdbook):
            """将xlrd的XF转换为xlwt的XFStyle"""
            xf = rdbook.xf_list[xf_index]
            font = rdbook.font_list[xf.font_index]
            
            # 创建字体
            wt_font = xlwt.Font()
            wt_font.name = font.name
            wt_font.height = font.height
            wt_font.bold = font.bold
            wt_font.italic = font.italic
            wt_font.underline = font.underline_type
            wt_font.colour_index = font.colour_index
            
            # 创建对齐
            wt_align = xlwt.Alignment()
            wt_align.horz = xf.alignment.hor_align
            wt_align.vert = xf.alignment.vert_align
            wt_align.wrap = xf.alignment.text_wrapped
            
            # 创建边框
            wt_borders = xlwt.Borders()
            wt_borders.left = xf.border.left_line_style
            wt_borders.right = xf.border.right_line_style
            wt_borders.top = xf.border.top_line_style
            wt_borders.bottom = xf.border.bottom_line_style
            wt_borders.left_colour = xf.border.left_colour_index
            wt_borders.right_colour = xf.border.right_colour_index
            wt_borders.top_colour = xf.border.top_colour_index
            wt_borders.bottom_colour = xf.border.bottom_colour_index
            
            # 创建图案（背景）
            wt_pattern = xlwt.Pattern()
            wt_pattern.pattern = xf.background.fill_pattern
            wt_pattern.pattern_fore_colour = xf.background.pattern_colour_index
            wt_pattern.pattern_back_colour = xf.background.background_colour_index
            
            # 创建保护
            wt_protection = xlwt.Protection()
            wt_protection.cell_locked = xf.protection.cell_locked
            wt_protection.formula_hidden = xf.protection.formula_hidden
            
            # 创建XFStyle
            style = xlwt.XFStyle()
            style.font = wt_font
            style.alignment = wt_align
            style.borders = wt_borders
            style.pattern = wt_pattern
            style.protection = wt_protection
            style.num_format_str = rdbook.format_map[xf.format_key].format_str if xf.format_key in rdbook.format_map else 'General'
            
            return style
        
        # 复制每个工作表
        for sheet_idx in range(rb.nsheets):
            rs = rb.sheet_by_index(sheet_idx)
            
            # 添加工作表
            ws = wb.add_sheet(rs.name, cell_overwrite_ok=True)
            
            # 复制列宽
            for col_idx in range(rs.ncols):
                try:
                    width = rb.sheet_by_index(sheet_idx).colinfo_map[col_idx].width
                    ws.col(col_idx).width = width
                except KeyError:
                    pass
            
            # 复制行高和单元格内容
            for row_idx in range(rs.nrows):
                try:
                    height = rb.sheet_by_index(sheet_idx).rowinfo_map[row_idx].height
                    ws.row(row_idx).height = height
                except KeyError:
                    pass
                
                for col_idx in range(rs.ncols):
                    cell = rs.cell(row_idx, col_idx)
                    cell_value = cell.value
                    xf_index = cell.xf_index
                    
                    # 替换变量
                    if cell_value and isinstance(cell_value, str):
                        new_value = self._replace_text(cell_value, data)
                    else:
                        new_value = cell_value
                    
                    # 获取样式
                    if xf_index < len(rb.xf_list):
                        style = get_xf_style(xf_index, rb)
                    else:
                        style = xlwt.Style.default_style
                    
                    # 写入单元格
                    ws.write(row_idx, col_idx, new_value, style)
        
        wb.save(self.output_path)
    
    def _extract_variables(self, text):
        """提取文本中的所有变量，支持嵌套的{}"""
        variables = []
        i = 0
        while i < len(text):
            if text[i] == '{':
                # 找到匹配的}
                count = 1
                j = i + 1
                while j < len(text) and count > 0:
                    if text[j] == '{':
                        count += 1
                    elif text[j] == '}':
                        count -= 1
                    j += 1
                
                if count == 0:
                    # 找到完整的变量
                    var_content = text[i+1:j-1]
                    variables.append((i, j, var_content))
                    i = j
                else:
                    i += 1
            else:
                i += 1
        
        return variables
    
    def _replace_text(self, text, data):
        """替换文本中的变量 - 支持嵌套的JavaScript模板字符串"""
        if not text or not isinstance(text, str):
            return text
        
        original_text = text
        
        # 提取所有变量（支持嵌套）
        variables = self._extract_variables(text)
        
        # 从后向前替换，避免位置变化
        for start, end, var_content in reversed(variables):
            var_clean = var_content.strip()
            # 同时尝试原始内容（保留空格）和清理后的内容
            if var_content in data:
                value = str(data[var_content])
                text = text[:start] + value + text[end:]
                print(f"替换: {{{var_content[:50]}...}} -> {value[:50]}...")
            elif var_clean in data:
                value = str(data[var_clean])
                text = text[:start] + value + text[end:]
                print(f"替换: {{{var_clean[:50]}...}} -> {value[:50]}...")
            elif var_clean.startswith('年月日_'):
                # 处理 {年月日_字段名} 格式，如 {年月日_handle_at} -> 2026年2月14日
                base_field = var_clean[4:]  # 去掉 '年月日_' 前缀
                if base_field in data:
                    date_value = data[base_field]
                    if date_value:
                        chinese_date = self._convert_to_chinese_date(date_value)
                        text = text[:start] + chinese_date + text[end:]
                        print(f"替换: {{{var_clean}}} -> {chinese_date}")
        
        return text
    
    def _convert_to_chinese_date(self, date_value):
        """将日期值转换为中文格式，如 2026-02-14 -> 2026年2月14日"""
        from datetime import datetime
        
        if not date_value:
            return ''
        
        date_str = str(date_value).strip()
        
        # 尝试解析各种日期格式
        formats = [
            '%Y-%m-%d',      # 2026-02-14
            '%Y/%m/%d',      # 2026/02/14
            '%Y%m%d',        # 20260214
            '%Y-%m-%d %H:%M:%S',  # 2026-02-14 10:30:00
            '%Y/%m/%d %H:%M:%S',  # 2026/02/14 10:30:00
        ]
        
        for fmt in formats:
            try:
                # 对于带时间的格式，只取日期部分
                if ' ' in date_str:
                    date_part = date_str.split()[0]
                else:
                    date_part = date_str
                
                # 尝试匹配格式
                if fmt in ['%Y-%m-%d', '%Y/%m/%d']:
                    dt = datetime.strptime(date_part, fmt)
                elif fmt == '%Y%m%d' and len(date_part) == 8:
                    dt = datetime.strptime(date_part, fmt)
                elif ' ' in date_str:
                    dt = datetime.strptime(date_str, fmt)
                else:
                    continue
                
                return f"{dt.year}年{dt.month}月{dt.day}日"
            except (ValueError, TypeError):
                continue
        
        # 如果都无法解析，返回原值
        return date_str
    
    def _replace_in_paragraph(self, para, data):
        """替换段落中的变量（Word）- 保留每个 run 的格式"""
        if not para.text:
            return
        
        # 检查段落中是否包含变量
        if '{' not in para.text or '}' not in para.text:
            return
        
        # 首先处理变量被分散在多个 runs 的情况
        # 合并包含同一变量的相邻 runs
        self._merge_runs_with_variables(para)
        
        # 现在每个变量应该都在一个单独的 run 中
        for run in para.runs:
            if '{' in run.text and '}' in run.text:
                original_text = run.text
                new_text = self._replace_text(original_text, data)
                if new_text != original_text:
                    run.text = new_text
                    # case_no 相关变量需要去除粗体
                    if '{case_no' in original_text or '{case_no_raw' in original_text:
                        run.bold = False
    
    def _merge_runs_with_variables(self, para):
        """合并包含同一变量的相邻 runs"""
        if len(para.runs) <= 1:
            return
        
        # 收集完整文本并找到变量位置
        full_text = ''.join(run.text or '' for run in para.runs)
        
        # 找出所有变量及其位置
        variables = self._extract_variables_with_positions(full_text)
        
        if not variables:
            return
        
        # 计算每个 run 的文本范围
        run_ranges = []
        pos = 0
        for run in para.runs:
            text_len = len(run.text or '')
            run_ranges.append((pos, pos + text_len, run))
            pos += text_len
        
        # 对于每个跨越多个 runs 的变量，合并这些 runs
        for start, end, var_content in sorted(variables, key=lambda x: x[0], reverse=True):
            # 找到包含这个变量的所有 runs
            runs_to_merge = []
            for run_start, run_end, run in run_ranges:
                if run_start < end and run_end > start:
                    runs_to_merge.append((run_start, run_end, run))
            
            if len(runs_to_merge) > 1:
                # 需要合并这些 runs
                # 使用第一个 run 的格式
                first_run = runs_to_merge[0][2]
                merged_text = ''.join(run.text or '' for _, _, run in runs_to_merge)
                
                # 设置第一个 run 的文本
                first_run.text = merged_text
                
                # 清空其他 runs
                for _, _, run in runs_to_merge[1:]:
                    run.text = ''
        
        # 清理空 runs（可选，保持段落整洁）
        # 注意：这里我们不能真正删除 runs，只能清空它们
    
    def _extract_variables_with_positions(self, text):
        """提取文本中的所有变量及其位置，支持嵌套的{}"""
        variables = []
        i = 0
        while i < len(text):
            if text[i] == '{':
                # 找到匹配的}
                count = 1
                j = i + 1
                while j < len(text) and count > 0:
                    if text[j] == '{':
                        count += 1
                    elif text[j] == '}':
                        count -= 1
                    j += 1
                
                if count == 0:
                    # 找到完整的变量
                    var_content = text[i+1:j-1]
                    variables.append((i, j, var_content))
                    i = j
                else:
                    i += 1
            else:
                i += 1
        
        return variables
    
    def _get_field(self, data, primary_key, fallback_keys=None, default=''):
        """
        从字典中获取字段值，支持多个备选键名
        
        Args:
            data: 字典数据
            primary_key: 主键名
            fallback_keys: 备选键名列表
            default: 默认值
        
        Returns:
            字段值或默认值
        """
        if not isinstance(data, dict):
            return default
        
        # 尝试主键
        if primary_key in data:
            value = data[primary_key]
            return value if value is not None else default
        
        # 尝试备选键
        if fallback_keys:
            for key in fallback_keys:
                if key in data:
                    value = data[key]
                    return value if value is not None else default
        
        return default
    
    def _preprocess_data(self, data, target_applicant=None, target_applicants=None):
        """
        预处理数据，计算所有变量值
        
        Args:
            data: 案件数据
            target_applicant: 目标申请人姓名（可选，单个）- 已废弃，请使用 target_applicants
            target_applicants: 目标申请人姓名列表（可选，支持多申请人）
        """
        result = {}
        
        # 兼容旧版本
        if target_applicant and not target_applicants:
            target_applicants = [target_applicant]
        
        # 获取案件数据（处理嵌套结构）
        # 尝试多种可能的数据结构
        case_data = data
        if 'data' in data:
            case_data = data['data']
            if isinstance(case_data, dict) and 'data' in case_data:
                case_data = case_data['data']
        
        print(f"案件数据 keys: {case_data.keys() if hasattr(case_data, 'keys') else 'N/A'}")
        print(f"目标申请人列表: {target_applicants}")
        
        # 1. 基础字段
        case_no_raw = case_data.get('case_no', '')
        # 如果 case_no 以"明"开头，则去掉"明"字
        if case_no_raw and case_no_raw.startswith('明'):
            case_no_raw = case_no_raw[1:]
        
        # 保存原始案号（保留方括号 [] 格式，用于立案审批表）
        result['case_no_raw'] = case_no_raw
        
        # 使用方括号 [] 格式
        case_no = case_no_raw
        result['case_no'] = case_no
        
        # 从案号中提取年份部分和编号，如 "永劳人仲案字[2026]123号"
        # case_no_year -> "永劳人仲案字[2026]"
        # case_no_no -> "123"
        case_no_year = ''
        case_no_no = ''
        if case_no_raw:
            # 提取年份部分（从开始到右方括号]，如：永劳人仲案字[2026]）
            year_part_match = re.search(r'^(.+?\[\d{4}\])', case_no_raw)
            if year_part_match:
                case_no_year = year_part_match.group(1)
            # 提取年后面的编号（右方括号后到"号"之前的数字）
            no_match = re.search(r'\](\d+)号', case_no_raw)
            if no_match:
                case_no_no = no_match.group(1)
        result['case_no_year'] = case_no_year
        result['case_no_no'] = case_no_no
        result['applicant'] = case_data.get('applicant', '')
        # 处理带空格的变量 { applicant}
        result[' applicant'] = case_data.get('applicant', '')
        result['applicant_str'] = case_data.get('applicant_str', '')
        result['respondent'] = case_data.get('respondent', '')
        result['apply_at'] = case_data.get('apply_at', '')
        result['handle_at'] = case_data.get('handle_at', '')
        result['case_reason'] = case_data.get('case_reason', '')
        
        # 当前年份
        result['year'] = str(datetime.now().year)
        
        # 2. 申请人信息
        applicant_arr = case_data.get('applicant_arr', [])
        print(f"申请人总数: {len(applicant_arr)}")
        
        # 如果指定了目标申请人列表，筛选出这些申请人
        selected_applicants = []
        if target_applicants and applicant_arr:
            for app in applicant_arr:
                app_name = app.get('name', '') or app.get('applicant_name', '')
                if app_name in target_applicants:
                    selected_applicants.append(app)
                    print(f"找到目标申请人: {app_name}")
            # 如果找不到任何匹配的申请人，使用第一个作为默认
            if not selected_applicants and applicant_arr:
                selected_applicants = [applicant_arr[0]]
                print(f"未找到匹配申请人，使用第一个: {selected_applicants[0].get('name', '')}")
        elif applicant_arr:
            # 没有指定目标申请人，使用所有申请人
            selected_applicants = applicant_arr
            print(f"未指定目标申请人，使用所有申请人")
        
        print(f"选中的申请人数量: {len(selected_applicants)}")
        
        # 填充申请人信息到 applicant_arr[i] 变量
        for idx, applicant in enumerate(selected_applicants):
            print(f"处理第 {idx+1} 个申请人: {applicant.get('name', '')}")
            
            # 基础信息
            result[f'applicant_arr[{idx}].name'] = self._get_field(applicant, 'name')
            result[f'applicant_arr[{idx}].mobile'] = self._get_field(applicant, 'mobile', ['phone', 'tel', 'telephone'])
            result[f'applicant_arr[{idx}].id_number'] = self._get_field(applicant, 'id_number')
            result[f'applicant_arr[{idx}].address'] = self._get_field(applicant, 'registered_permanent_residence', ['address'])
            result[f'applicant_arr[{idx}].registered_permanent_residence'] = self._get_field(applicant, 'registered_permanent_residence')
            
            # 申请人代理人信息
            agents = applicant.get('agents', [])
            if agents and len(agents) > 0:
                first_agent = agents[0]
                result[f'applicant_arr[{idx}].agents[0].name'] = self._get_field(first_agent, 'name')
                result[f'applicant_arr[{idx}].agents[0].mobile'] = self._get_field(first_agent, 'mobile', ['phone', 'tel', 'telephone'])
            else:
                result[f'applicant_arr[{idx}].agents[0].name'] = ''
                result[f'applicant_arr[{idx}].agents[0].mobile'] = ''
        
        # 兼容旧版本：如果没有选中任何申请人，填充空值
        if not selected_applicants:
            result['applicant_arr[0].name'] = ''
            result['applicant_arr[0].mobile'] = ''
            result['applicant_arr[0].id_number'] = ''
            result['applicant_arr[0].address'] = ''
            result['applicant_arr[0].registered_permanent_residence'] = ''
            result['applicant_arr[0].agents[0].name'] = ''
            result['applicant_arr[0].agents[0].mobile'] = ''
            result['registered_permanent_residence'] = ''
            result['applicant_name'] = ''
            result['applicant_mobile'] = ''
            result['applicant_id_number'] = ''
            result['applicant_agent_name'] = ''
            result['applicant_agent_mobile'] = ''
            print("无申请人信息")
        else:
            # 第一个申请人的信息也填充到旧版本变量中（兼容）
            first_applicant = selected_applicants[0]
            result['registered_permanent_residence'] = self._get_field(first_applicant, 'registered_permanent_residence')
            result['applicant_name'] = self._get_field(first_applicant, 'name')
            result['applicant_mobile'] = self._get_field(first_applicant, 'mobile', ['phone', 'tel', 'telephone'])
            result['applicant_id_number'] = self._get_field(first_applicant, 'id_number')
            
            agents = first_applicant.get('agents', [])
            if agents and len(agents) > 0:
                first_agent = agents[0]
                result['applicant_agent_name'] = self._get_field(first_agent, 'name')
                result['applicant_agent_mobile'] = self._get_field(first_agent, 'mobile', ['phone', 'tel', 'telephone'])
            else:
                result['applicant_agent_name'] = ''
                result['applicant_agent_mobile'] = ''
        
        # 2.1 生成选中申请人的信息字符串 {a_str}
        # 格式: 申请人N：姓名，性别，民族，出生日期出生，身份证住址。公民身份号码：身份证号。
        # 注意：除第一个申请人外，前面加两个全角空格模拟首行缩进
        if selected_applicants:
            applicant_str_parts = []
            for idx, applicant in enumerate(selected_applicants):
                name = self._get_field(applicant, 'name')
                sex = self._get_field(applicant, 'sex')
                nation = self._get_field(applicant, 'nation')
                birth = self._get_field(applicant, 'birth')
                residence = self._get_field(applicant, 'registered_permanent_residence')
                id_number = self._get_field(applicant, 'id_number')
                
                # 格式化出生日期: 2004-07-08 -> 2004年7月8日
                birth_formatted = birth
                if birth and len(birth) >= 10:
                    try:
                        birth_dt = datetime.strptime(birth.split()[0], '%Y-%m-%d')
                        birth_formatted = f"{birth_dt.year}年{birth_dt.month}月{birth_dt.day}日"
                    except:
                        birth_formatted = birth
                
                # 构建单个申请人信息字符串
                # 除第一个申请人外，前面加两个全角空格（　）模拟首行缩进
                indent = '' if idx == 0 else '　　'
                part = f"{indent}申请人{idx + 1}：{name}，{sex}，{nation}，{birth_formatted}出生，身份证住址：{residence}。公民身份号码：{id_number}。"
                applicant_str_parts.append(part)
            
            result['a_str'] = '\n'.join(applicant_str_parts)
        else:
            result['a_str'] = ''
        
        # 3. 被申请人信息（支持多个）
        respondent_arr = case_data.get('respondent_arr', [])
        print(f"被申请人数量: {len(respondent_arr)}")
        
        # 第一个被申请人
        if respondent_arr and len(respondent_arr) > 0:
            first_respondent = respondent_arr[0]
            print(f"第一个被申请人 keys: {list(first_respondent.keys()) if isinstance(first_respondent, dict) else 'N/A'}")
            
            result['respondent_arr[0].name'] = self._get_field(first_respondent, 'name')
            result['respondent_arr[0].company_address'] = self._get_field(first_respondent, 'company_address', ['address'])
            result['respondent_name'] = self._get_field(first_respondent, 'name')
            result['respondent_address'] = self._get_field(first_respondent, 'company_address', ['address'])
            result['respondent_social_code'] = self._get_field(first_respondent, 'social_code')
            result['respondent_legal_name'] = self._get_field(first_respondent, 'legal_name')
            result['respondent_arr[0].legal_name'] = self._get_field(first_respondent, 'legal_name')
            # 支持多种可能的字段名：legal_mobile, legal_phone, legal_tel, phone, mobile
            result['respondent_arr[0].legal_mobile'] = self._get_field(
                first_respondent, 'legal_mobile', 
                ['legal_phone', 'legal_tel', 'phone', 'mobile', 'tel', 'telephone']
            )
            
            print(f"respondent_arr[0].legal_mobile 赋值: '{result['respondent_arr[0].legal_mobile']}'")
            
            # 被申请人代理人信息
            agents = first_respondent.get('agents', [])
            print(f"第一个被申请人代理人数量: {len(agents)}")
            
            if agents and len(agents) > 0:
                first_agent = agents[0]
                print(f"第一个被申请人第一个代理人 keys: {list(first_agent.keys()) if isinstance(first_agent, dict) else 'N/A'}")
                
                result['respondent_arr[0].agents[0].name'] = self._get_field(first_agent, 'name')
                result['respondent_arr[0].agents[0].mobile'] = self._get_field(
                    first_agent, 'mobile', 
                    ['phone', 'tel', 'telephone']
                )
                result['respondent_agent_name'] = self._get_field(first_agent, 'name')
                result['respondent_agent_mobile'] = self._get_field(
                    first_agent, 'mobile', 
                    ['phone', 'tel', 'telephone']
                )
                
                print(f"respondent_arr[0].agents[0].mobile 赋值: '{result['respondent_arr[0].agents[0].mobile']}'")
            else:
                result['respondent_arr[0].agents[0].name'] = ''
                result['respondent_arr[0].agents[0].mobile'] = ''
                result['respondent_agent_name'] = ''
                result['respondent_agent_mobile'] = ''
                print("respondent_arr[0].agents[0].mobile 赋值为空字符串（无代理人）")
        else:
            result['respondent_arr[0].name'] = ''
            result['respondent_arr[0].company_address'] = ''
            result['respondent_name'] = ''
            result['respondent_address'] = ''
            result['respondent_social_code'] = ''
            result['respondent_legal_name'] = ''
            result['respondent_arr[0].legal_name'] = ''
            result['respondent_arr[0].legal_mobile'] = ''
            result['respondent_arr[0].agents[0].name'] = ''
            result['respondent_arr[0].agents[0].mobile'] = ''
            result['respondent_agent_name'] = ''
            result['respondent_agent_mobile'] = ''
            print("respondent_arr[0].legal_mobile 和 respondent_arr[0].agents[0].mobile 赋值为空字符串（无被申请人）")
        
        # 3.1 生成被申请人信息字符串 {r_str}
        # 格式: 
        # 被申请人：XXX，统一社会信用代码: XXX，住所：XXX。
        # 　　法定代表人：XXX，职务。
        if respondent_arr and len(respondent_arr) > 0:
            first_respondent = respondent_arr[0]
            r_name = self._get_field(first_respondent, 'name')
            r_social_code = self._get_field(first_respondent, 'social_code')
            r_address = self._get_field(first_respondent, 'company_address', ['address'])
            r_legal_name = self._get_field(first_respondent, 'legal_name')
            # 尝试获取职务，可能字段名为 duty, position, legal_duty 等
            r_duty = self._get_field(first_respondent, 'duty', ['position', 'legal_duty', 'legal_position', 'job', 'title'])
            
            # 第一行：被申请人基本信息
            line1 = f"被申请人：{r_name}，统一社会信用代码: {r_social_code}，住所：{r_address}。"
            # 第二行：法定代表人信息（前面加两个全角空格模拟缩进）
            if r_duty:
                line2 = f"　　法定代表人：{r_legal_name}，{r_duty}。"
            else:
                line2 = f"　　法定代表人：{r_legal_name}。"
            
            result['r_str'] = f"{line1}\n{line2}"
        else:
            result['r_str'] = ''
        
        # 第二个被申请人（如果存在）
        if respondent_arr and len(respondent_arr) > 1:
            second_respondent = respondent_arr[1]
            result['respondent_arr[1].name'] = self._get_field(second_respondent, 'name')
            result['respondent_arr[1].legal_name'] = self._get_field(second_respondent, 'legal_name')
            result['respondent_arr[1].legal_mobile'] = self._get_field(
                second_respondent, 'legal_mobile',
                ['legal_phone', 'legal_tel', 'phone', 'mobile', 'tel', 'telephone']
            )
            
            agents = second_respondent.get('agents', [])
            if agents and len(agents) > 0:
                first_agent = agents[0]
                result['respondent_arr[1].agents[0].name'] = self._get_field(first_agent, 'name')
                result['respondent_arr[1].agents[0].mobile'] = self._get_field(
                    first_agent, 'mobile',
                    ['phone', 'tel', 'telephone']
                )
            else:
                result['respondent_arr[1].agents[0].name'] = ''
                result['respondent_arr[1].agents[0].mobile'] = ''
        else:
            # 确保第二个被申请人的变量也有默认值
            result['respondent_arr[1].name'] = ''
            result['respondent_arr[1].legal_name'] = ''
            result['respondent_arr[1].legal_mobile'] = ''
            result['respondent_arr[1].agents[0].name'] = ''
            result['respondent_arr[1].agents[0].mobile'] = ''
        
        # 4. 仲裁请求列表
        case_arb_request = case_data.get('case_arb_request', [])
        
        # 请求数量
        result['request_count'] = str(len(case_arb_request))
        
        if case_arb_request:
            # 生成请求编号列表 "1、2、3"
            request_numbers = '、'.join([str(i + 1) for i in range(len(case_arb_request))])
            result['request_numbers'] = request_numbers
            
            # 每个请求的 intro 和 object
            for i, req in enumerate(case_arb_request):
                result[f'request_{i+1}_intro'] = req.get('intro', '')
                result[f'request_{i+1}_object'] = req.get('object', '')
                # 支持 {item.intro} 形式的变量（在表格循环中）
                if i == 0:
                    result['item.intro'] = req.get('intro', '')
            
            # 生成请求列表文本（用于显示）
            request_list = []
            for i, req in enumerate(case_arb_request):
                intro = req.get('intro', '')
                if intro:
                    request_list.append(f"{i+1}. {intro}")
            result['request_list'] = '\n'.join(request_list)
            
            # 生成JavaScript风格的表格循环内容
            # 完整的表达式: data.case_arb_request.map((item, idx) => `${idx + 1}. ${item.intro}`).join('\n')
            request_lines = []
            for i, req in enumerate(case_arb_request):
                intro = req.get('intro', '')
                request_lines.append(f"{i+1}. {intro}")
            request_text = '\n'.join(request_lines)
            
            # 支持这种复杂的模板表达式 - 注意转义字符的处理
            result["data.case_arb_request.map((item, idx) => `${idx + 1}. ${item.intro}`).join('\\n')"] = request_text
        else:
            result['request_numbers'] = ''
            result['request_list'] = ''
            result['item.intro'] = ''
            result["data.case_arb_request.map((item, idx) => `${idx + 1}. ${item.intro}`).join('\\n')"] = ''
        
        # 5. 案件标的总和
        total_money = sum(
            float(req.get('object', 0) or 0) 
            for req in case_arb_request
        )
        result['total_money'] = f"{total_money:.2f}"
        # 支持 JavaScript 表达式形式的变量名
        result['case_arb_request.reduce((sum, item) => sum + parseFloat(item.object || 0), 0).toFixed(2)'] = f"{total_money:.2f}"
        
        # 6. 第三人信息
        thirdpartys = case_data.get('thirdpartys', '')
        result['thirdpartys'] = thirdpartys if thirdpartys else ''
        
        # 7. 中文日期（今日）
        result['中文_today'] = self._get_chinese_date()
        result['today'] = datetime.now().strftime('%Y-%m-%d')
        
        # 8. 立案日期格式化（中文）
        handle_at = case_data.get('handle_at', '')
        if handle_at:
            try:
                dt = datetime.strptime(handle_at.split()[0], '%Y-%m-%d')
                result['handle_at_chinese'] = self._get_chinese_date(dt)
                result['中文_handle_at'] = self._get_chinese_date(dt)
                # 年月日格式 (2026年3月5日)
                result['年月日_handle_at'] = f"{dt.year}年{dt.month}月{dt.day}日"
                # 立案日期年月日单独字段
                result['handle_at_y'] = str(dt.year)  # 年，如 2026
                result['handle_at_m'] = str(dt.month)  # 月 1-12，无前导0
                result['handle_at_d'] = str(dt.day)    # 日 1-31，无前导0
            except:
                result['handle_at_chinese'] = handle_at
                result['中文_handle_at'] = handle_at
                result['年月日_handle_at'] = handle_at
                result['handle_at_y'] = ''
                result['handle_at_m'] = ''
                result['handle_at_d'] = ''
        else:
            result['handle_at_chinese'] = ''
            result['中文_handle_at'] = ''
            result['年月日_handle_at'] = ''
            result['handle_at_y'] = ''
            result['handle_at_m'] = ''
            result['handle_at_d'] = ''
        
        # 8.1 申请日期格式化
        apply_at = case_data.get('apply_at', '')
        if apply_at:
            try:
                dt = datetime.strptime(apply_at.split()[0], '%Y-%m-%d')
                result['中文_apply_at'] = self._get_chinese_date(dt)
                # 年月日格式 (2026年3月1日)
                result['年月日_apply_at'] = f"{dt.year}年{dt.month}月{dt.day}日"
            except:
                result['中文_apply_at'] = apply_at
                result['年月日_apply_at'] = apply_at
        else:
            result['中文_apply_at'] = ''
            result['年月日_apply_at'] = ''
        
        # 9. 开庭信息（从 tribunal_plan 最后一个元素提取）
        result['open'] = self._get_open_date_time(case_data)
        
        # 10. 开庭电话（从 tribunal_plan 最后一个元素提取 tel）
        result['tel'] = self._get_tribunal_tel(case_data)
        
        # 11. 开庭创建日期（从 tribunal_plan 最后一个元素提取 created_at，格式化为中文日期）
        tribunal_plan = case_data.get('tribunal_plan', [])
        if tribunal_plan and isinstance(tribunal_plan, list):
            last_plan = tribunal_plan[-1]
            if isinstance(last_plan, dict):
                create_at = last_plan.get('created_at', '')
                if create_at:
                    try:
                        dt = datetime.strptime(create_at.split()[0], '%Y-%m-%d')
                        chinese_date = self._get_chinese_date(dt)
                        result['年月日_create_at'] = chinese_date
                        result['年月日_created_at'] = chinese_date  # 兼容带 d 的写法
                    except:
                        result['年月日_create_at'] = create_at
                        result['年月日_created_at'] = create_at  # 兼容带 d 的写法
                else:
                    result['年月日_create_at'] = ''
                    result['年月日_created_at'] = ''  # 兼容带 d 的写法
            else:
                result['年月日_create_at'] = ''
                result['年月日_created_at'] = ''  # 兼容带 d 的写法
        else:
            result['年月日_create_at'] = ''
            result['年月日_created_at'] = ''  # 兼容带 d 的写法
        
        # 12. 仲裁员和书记员信息
        result['arbitrator'] = case_data.get('arbitrator', '') or ''
        result['arbitrator_one'] = case_data.get('arbitrator_one', '') or ''
        result['arbitrator_two'] = case_data.get('arbitrator_two', '') or ''
        result['clerk'] = case_data.get('clerk', '') or ''
        
        # 13. 反申请日期（从 review_matter 中查找 apply_matter='仲裁反申请' 的条目，取 start_at）
        re_at = self._get_reverse_request_date(case_data)
        if re_at:
            try:
                # 解析日期，格式可能是 "2025-12-31 00:00:00"
                dt = datetime.strptime(re_at.split()[0], '%Y-%m-%d')
                result['re_at_y'] = str(dt.year)  # 年，如 2026
                result['re_at_m'] = str(dt.month)  # 月 1-12，无前导0
                result['re_at_d'] = str(dt.day)    # 日 1-31，无前导0
                result['中文_re_at'] = self._get_chinese_date(dt)  # 中文日期，如 二零二五年十二月三十一日
            except:
                result['re_at_y'] = ''
                result['re_at_m'] = ''
                result['re_at_d'] = ''
                result['中文_re_at'] = ''
        else:
            result['re_at_y'] = ''
            result['re_at_m'] = ''
            result['re_at_d'] = ''
            result['中文_re_at'] = ''
        
        return result
    
    def _get_chinese_date(self, date=None):
        """获取中文日期格式：二〇二六年一月一日"""
        if date is None:
            date = datetime.now()
        elif isinstance(date, str):
            date = datetime.strptime(date.split()[0], '%Y-%m-%d')
        
        chinese_nums = {
            '0': '〇', '1': '一', '2': '二', '3': '三', '4': '四',
            '5': '五', '6': '六', '7': '七', '8': '八', '9': '九'
        }
        
        # 年份：直接替换数字
        year_str = ''.join(chinese_nums.get(c, c) for c in str(date.year))
        
        # 月份转换
        month = date.month
        month_map = {
            1: '一', 2: '二', 3: '三', 4: '四', 5: '五',
            6: '六', 7: '七', 8: '八', 9: '九', 10: '十',
            11: '十一', 12: '十二'
        }
        month_str = month_map.get(month, str(month))
        
        # 日期转换
        day = date.day
        if day <= 10:
            day_map = {1: '一', 2: '二', 3: '三', 4: '四', 5: '五',
                      6: '六', 7: '七', 8: '八', 9: '九', 10: '十'}
            day_str = day_map.get(day, str(day))
        elif day < 20:
            # 11-19: 十一、十二...
            day_map = {11: '十一', 12: '十二', 13: '十三', 14: '十四', 15: '十五',
                      16: '十六', 17: '十七', 18: '十八', 19: '十九'}
            day_str = day_map.get(day, f'十{day - 10}')
        elif day == 20:
            day_str = '二十'
        elif day < 30:
            # 21-29
            day_map = {21: '二十一', 22: '二十二', 23: '二十三', 24: '二十四', 
                      25: '二十五', 26: '二十六', 27: '二十七', 28: '二十八', 
                      29: '二十九'}
            day_str = day_map.get(day, f'二十{day - 20}')
        elif day == 30:
            day_str = '三十'
        else:
            day_str = '三十一'
        
        return f"{year_str}年{month_str}月{day_str}日"
    
    def _get_open_date_time(self, case_data):
        """
        获取开庭日期时间格式：2026年3月11日（星期三）上午9时
        从 tribunal_plan 最后一个元素提取 open_at 和 text 字段
        """
        tribunal_plan = case_data.get('tribunal_plan', [])
        if not tribunal_plan or not isinstance(tribunal_plan, list):
            return ''
        
        # 获取最后一个元素
        last_plan = tribunal_plan[-1]
        if not isinstance(last_plan, dict):
            return ''
        
        open_at = last_plan.get('open_at', '')
        text = last_plan.get('text', '')
        
        if not open_at:
            return ''
        
        try:
            # 解析 open_at (如: 2026-03-11)
            dt = datetime.strptime(open_at.split()[0], '%Y-%m-%d')
            
            # 获取星期几
            weekdays = ['星期一', '星期二', '星期三', '星期四', '星期五', '星期六', '星期日']
            weekday = weekdays[dt.weekday()]
            
            # 从 text 提取小时 (如: "3月11日09时第1次公开开庭" -> 09)
            hour = None
            if text:
                # 匹配 "XX时" 格式
                hour_match = re.search(r'(\d{1,2})时', text)
                if hour_match:
                    hour = int(hour_match.group(1))
            
            # 格式化时间描述
            if hour is not None:
                # 判断上午/下午
                if hour < 12:
                    am_pm = '上午'
                elif hour == 12:
                    am_pm = '中午'
                else:
                    am_pm = '下午'
                    hour = hour - 12 if hour > 12 else hour
                time_str = f"{am_pm}{hour}时"
            else:
                time_str = ''
            
            # 组合结果: 2026年3月11日（星期三）上午9时
            if time_str:
                return f"{dt.year}年{dt.month}月{dt.day}日（{weekday}）{time_str}"
            else:
                return f"{dt.year}年{dt.month}月{dt.day}日（{weekday}）"
        except Exception as e:
            print(f"解析开庭日期失败: {e}, open_at={open_at}, text={text}")
            return open_at
    
    def _get_tribunal_tel(self, case_data):
        """
        获取开庭电话（从 tribunal_plan 最后一个元素提取 tel 字段）
        如果获取不到则返回默认值：0598-3650311
        """
        default_tel = '0598-3650311'
        
        tribunal_plan = case_data.get('tribunal_plan', [])
        if not tribunal_plan or not isinstance(tribunal_plan, list):
            return default_tel
        
        # 获取最后一个元素
        last_plan = tribunal_plan[-1]
        if not isinstance(last_plan, dict):
            return default_tel
        
        tel = last_plan.get('tel', '')
        return tel if tel else default_tel
    
    def _get_reverse_request_date(self, case_data):
        """
        获取反申请日期（从 review_matter 中查找 apply_matter 包含'反申请'字样的条目，取 start_at 字段）
        这是反申请的实际提交日期，比 created_at 更准确
        """
        review_matter = case_data.get('review_matter', [])
        if not review_matter or not isinstance(review_matter, list):
            return ''
        
        # 查找 apply_matter 包含'反申请'字样的条目
        for matter in review_matter:
            if isinstance(matter, dict):
                apply_matter = matter.get('apply_matter', '') or ''
                if '反申请' in apply_matter:
                    start_at = matter.get('start_at', '')
                    if start_at:
                        return start_at
        
        return ''


def generate_document(template_name, output_name, case_id):
    """
    生成单个文档
    """
    import requests
    
    # 获取案件详情
    response = requests.get(f'http://localhost:5000/api/handle/detail?id={case_id}')
    data = response.json()
    
    # 生成文档
    template_path = f'文件生成/1-立案/{template_name}'
    output_path = f'文件生成/1-立案/output/{output_name}'
    
    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    generator = DocumentGenerator(template_path, output_path)
    generator.generate(data)
    
    print(f"文档已生成: {output_path}")
    return output_path


if __name__ == '__main__':
    # 测试生成
    # generate_document('（受理）立案审批表.docx', '测试输出.docx', '199966')
    
    # 测试中文日期
    gen = DocumentGenerator('文件生成/1-立案/变更仲裁申请受理通知书（申请人）.docx', '/tmp/test.docx')
    print("中文日期测试:", gen._get_chinese_date())
