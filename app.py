#!/usr/bin/env python3
"""
劳动仲裁信息查询综合服务平台
整合功能：
1. 前端页面服务（劳动仲裁申请书填写、案件查询）
2. 内部API服务（企业信息查询、身份证信息查询）
3. 数据库API服务（案件数据增删改查）

部署在 10.99.144.x 网段
"""

import sys
import os
import io

# 添加modules目录到Python路径
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'modules'))

# 导入笔录提取模块
from court_record_extractor import extract_court_record, format_date_no_leading_zero


def format_handle_at(handle_at: str) -> str:
    """将 handle_at 格式化为 XXXX年X月X日（月份日期不带前导零）"""
    if not handle_at:
        return ""
    match = re.search(r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})', str(handle_at))
    if match:
        year = match.group(1)
        month = int(match.group(2))
        day = int(match.group(3))
        return f"{year}年{month}月{day}日"
    return handle_at

from flask import Flask, request, jsonify, render_template, send_from_directory
from flask_cors import CORS
import logging
from datetime import datetime
import pymysql
import json
import requests
import re
from docx.shared import Pt
from docx.oxml.ns import qn

# 导入原有模块
from config import Config
from login_manager import get_login_manager
from company_query import get_company_query
from id_card_query import get_id_card_query_manager

# 休假日缓存 { '202602': [23, 24, ...] }
_holiday_cache = {}
_holiday_cache_month = None
from database import get_db_manager

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# 创建Flask应用
app = Flask(__name__, static_folder='static', template_folder='templates')
app.config['TEMPLATES_AUTO_RELOAD'] = True  # 启用模板自动重新加载
CORS(app)  # 启用CORS支持

# 初始化管理器
login_manager = get_login_manager()
company_query = get_company_query()
db_manager = get_db_manager()

# 添加全局缓存控制头
@app.after_request
def add_cache_control_headers(response):
    """为所有响应添加缓存控制头"""
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response

# ============================================
# 页面路由
# ============================================

@app.route('/')
def index():
    """首页 - 劳动仲裁申请书在线填写"""
    return render_template('index.html')


@app.route('/query')
def query_page():
    """案件查询页面"""
    return render_template('query_test.html')


@app.route('/cases')
def cases_manage_page():
    """案件管理页面"""
    return render_template('cases_manage.html')


@app.route('/receive_query')
def receive_query_page():
    """收件查询页面"""
    return render_template('receive_query.html')


@app.route('/receive_detail')
def receive_detail_page():
    """收件详情页面"""
    return render_template('receive_detail.html')


@app.route('/handle_query')
def handle_query_page():
    """立案查询页面"""
    return render_template('handle_query.html')


@app.route('/handle_detail')
def handle_detail_page():
    """立案详情页面"""
    return render_template('handle_detail.html')


@app.route('/award/make')
def award_make_page():
    """裁决书制作页面"""
    # 根据客户端IP自动选择Dify地址
    client_ip = request.remote_addr
    
    # 判断客户端所属网段
    if client_ip.startswith('10.99.144.'):
        # 10.99.144.x 网段使用 10.99.144.29
        dify_url = 'http://10.99.144.29:8020'
    else:
        # 其他网段（包括192.168.123.x）使用 192.168.123.16
        dify_url = 'http://192.168.123.16:8020'
    
    return render_template('award_make.html', dify_url=dify_url)


@app.route('/reserve_query')
def reserve_query_page():
    """预约仲裁申请查询页面"""
    return render_template('reserve_query.html')


@app.route('/reserve_detail')
def reserve_detail_page():
    """预约仲裁申请详情页面"""
    return render_template('reserve_detail.html')


@app.route('/files/empty.docx')
def serve_empty_docx():
    """提供 empty.docx 文件下载，用于一键生成初稿时无可用文件的默认选项"""
    return send_from_directory('empty_word', 'empty.docx', as_attachment=False)


@app.route('/files/evidence/<path:filename>')
def serve_evidence_file(filename):
    """提供证据文件下载，用于Dify访问"""
    evidence_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'evidence_files')
    return send_from_directory(evidence_dir, filename, as_attachment=False)


# ============================================
# 内部API服务（原Info项目）
# ============================================

@app.route('/api/status', methods=['GET'])
def api_status():
    """API状态检查"""
    return jsonify({
        'code': 200,
        'message': 'API服务运行正常',
        'service': '劳动仲裁信息查询综合服务平台',
        'timestamp': datetime.now().isoformat(),
        'status': {
            'flask': 'running',
            'login_service': 'connected' if login_manager else 'disconnected',
            'query_service': 'connected' if company_query else 'disconnected',
            'database': 'connected' if db_manager.pool else 'disconnected'
        }
    })


@app.route('/api/login/status', methods=['GET'])
def login_status():
    """获取登录状态"""
    status = login_manager.get_login_status()
    return jsonify({
        'code': 200,
        'message': '登录状态查询成功',
        'data': status,
        'timestamp': datetime.now().isoformat()
    })


@app.route('/api/login', methods=['POST'])
def manual_login():
    """手动登录接口"""
    try:
        data = request.get_json()
        force = data.get('force', False) if data else False
        
        login_result = login_manager.login(force=force)
        
        return jsonify({
            'code': login_result['code'],
            'message': login_result['message'],
            'data': {
                'authKey': login_result.get('authKey'),
                'sessionId': login_result.get('sessionId'),
                'expiry_time': login_result.get('expiry_time')
            },
            'timestamp': datetime.now().isoformat()
        })
    except Exception as e:
        logger.error(f"手动登录接口错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器错误: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


@app.route('/api/company/query', methods=['POST'])
def query_company():
    """查询企业信息接口"""
    try:
        data = request.get_json()
        
        if not data or 'company_name' not in data:
            return jsonify({
                'code': 400,
                'message': '缺少参数: company_name',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 400
        
        company_name = data['company_name']
        exact_match = data.get('exact_match', True)
        
        logger.info(f"API查询请求: {company_name}, 精确匹配: {exact_match}")
        
        query_result = company_query.query_company_info(company_name, exact_match)
        
        response_data = {
            'code': query_result['code'],
            'message': query_result['message'],
            'query_info': {
                'company_name': company_name,
                'exact_match': exact_match,
                'total_count': query_result.get('total_count', 0),
                'matched_count': query_result.get('matched_count', 0),
                'source': query_result.get('source', 'unknown')
            },
            'timestamp': datetime.now().isoformat()
        }
        
        if query_result.get('data'):
            format_data = data.get('format', True)
            if format_data:
                formatted_data = []
                for item in query_result['data']:
                    formatted_item = company_query.format_company_info(item)
                    if formatted_item:
                        formatted_data.append(formatted_item)
                response_data['data'] = formatted_data
            else:
                response_data['data'] = query_result['data']
        else:
            response_data['data'] = None
        
        return jsonify(response_data)
        
    except Exception as e:
        logger.error(f"查询企业信息接口错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器错误: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


@app.route('/api/idcard/query', methods=['POST'])
def query_id_card():
    """查询身份证信息接口"""
    try:
        data = request.get_json()
        
        if not data or 'AAC147' not in data:
            return jsonify({
                'code': 400,
                'message': '缺少参数: AAC147 (身份证号码)',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 400
        
        id_card_number = data['AAC147']
        
        logger.info(f"身份证查询请求: {id_card_number}")
        
        id_card_query_manager = get_id_card_query_manager()
        query_result = id_card_query_manager.query_id_card(id_card_number)
        
        response_data = {
            'code': query_result['code'],
            'message': query_result['message'],
            'data': query_result.get('data'),
            'timestamp': datetime.now().isoformat()
        }
        
        return jsonify(response_data)
        
    except Exception as e:
        logger.error(f"查询身份证信息接口错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器错误: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


@app.route('/api/datashare/query', methods=['POST'])
def datashare_query():
    """
    数据共享查询接口
    支持三种查询类型：
    1. idcard - 身份证信息查询 (参数: id_number)
    2. company - 企业信息查询 (参数: company_name 或 credit_code)
    3. office - 机关单位信息查询 (参数: office_name)
    """
    try:
        data = request.get_json()
        
        if not data or 'type' not in data:
            return jsonify({
                'code': 400,
                'message': '缺少参数: type (查询类型: idcard/company/office)',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 400
        
        query_type = data['type']
        
        # 检查并续期登录状态
        if not login_manager.check_and_renew_login():
            return jsonify({
                'code': 401,
                'message': '登录失败，无法查询数据',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 401
        
        headers = login_manager.get_auth_headers()
        
        if query_type == 'idcard':
            # 身份证信息查询
            id_number = data.get('id_number')
            if not id_number:
                return jsonify({
                    'code': 400,
                    'message': '缺少参数: id_number (身份证号码)',
                    'data': None,
                    'timestamp': datetime.now().isoformat()
                }), 400
            
            logger.info(f"数据共享-身份证查询: {id_number}")
            
            # 调用身份证查询接口
            id_card_query_manager = get_id_card_query_manager()
            query_result = id_card_query_manager.query_id_card(id_number)
            
            return jsonify({
                'code': query_result['code'],
                'message': query_result['message'],
                'data': query_result.get('data'),
                'source': query_result.get('source', 'api'),
                'timestamp': datetime.now().isoformat()
            })
            
        elif query_type == 'company':
            # 企业信息查询
            company_name = data.get('company_name')
            credit_code = data.get('credit_code')
            
            if not company_name and not credit_code:
                return jsonify({
                    'code': 400,
                    'message': '缺少参数: company_name 或 credit_code',
                    'data': None,
                    'timestamp': datetime.now().isoformat()
                }), 400
            
            query_key = company_name or credit_code
            logger.info(f"数据共享-企业查询: {query_key}")
            
            # 调用企业查询接口
            company_query = get_company_query()
            
            if credit_code:
                # 按统一社会信用代码查询
                query_result = company_query.query_company_by_credit_code(credit_code)
            else:
                # 按企业名称查询
                query_result = company_query.query_company_info(company_name, exact_match=True)
            
            return jsonify({
                'code': query_result['code'],
                'message': query_result['message'],
                'data': query_result.get('data'),
                'source': query_result.get('source', 'api'),
                'timestamp': datetime.now().isoformat()
            })
            
        elif query_type == 'office':
            # 机关单位信息查询
            office_name = data.get('office_name')
            if not office_name:
                return jsonify({
                    'code': 400,
                    'message': '缺少参数: office_name (机关单位名称)',
                    'data': None,
                    'timestamp': datetime.now().isoformat()
                }), 400
            
            logger.info(f"数据共享-机关单位查询: {office_name}")
            
            # 调用机关单位查询接口
            url = "http://10.96.10.78:8080/v1/api/admin/datashare/openPlatformProxy/api/SJCK_officeUnitInfo"
            payload = {"JGMC": office_name}
            
            response = requests.post(url, headers=headers, json=payload, timeout=30)
            response.raise_for_status()
            result = response.json()
            
            if result.get('code') == 200:
                return jsonify({
                    'code': 200,
                    'message': '查询成功',
                    'data': result.get('data'),
                    'source': 'api',
                    'timestamp': datetime.now().isoformat()
                })
            else:
                return jsonify({
                    'code': result.get('code', 500),
                    'message': result.get('message', '查询失败'),
                    'data': None,
                    'timestamp': datetime.now().isoformat()
                })
        else:
            return jsonify({
                'code': 400,
                'message': f'不支持的查询类型: {query_type}',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 400
            
    except requests.exceptions.RequestException as e:
        logger.error(f"数据共享查询请求失败: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'查询请求失败: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500
    except Exception as e:
        logger.error(f"数据共享查询接口错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器错误: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


@app.route('/api/db/status', methods=['GET'])
def db_status():
    """获取数据库状态"""
    try:
        db_connected = db_manager.pool is not None
        login_records = db_manager.get_all_logins()
        expired_deleted = db_manager.delete_expired_logins()
        
        status_info = {
            'database_connected': db_connected,
            'total_login_records': len(login_records),
            'expired_records_deleted': expired_deleted,
            'login_records': login_records,
            'timestamp': datetime.now().isoformat()
        }
        
        return jsonify({
            'code': 200,
            'message': '数据库状态查询成功',
            'data': status_info,
            'timestamp': datetime.now().isoformat()
        })
    except Exception as e:
        logger.error(f"数据库状态查询错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'数据库状态查询错误: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


# ============================================
# 数据库API服务（原db_api_server.py）
# ============================================

def get_labor_db_connection():
    """获取劳动仲裁数据库连接"""
    return pymysql.connect(
        host=Config.DB_HOST,
        port=Config.DB_PORT,
        user=Config.DB_USER,
        password=Config.DB_PASSWORD,
        database=Config.DB_NAME,
        charset='utf8mb4'
    )


@app.route('/api/cases/save', methods=['POST'])
def save_case():
    """保存完整案件数据（新建或更新）"""
    data = request.json
    receipt_number = data.get('receipt_number')
    applicants = data.get('applicants', [])
    respondents = data.get('respondents', [])
    evidence = data.get('evidence', [])
    case_id = data.get('case_id')  # 编辑模式时传入
    mode = data.get('mode', 'create')  # 'create' 或 'update'
    
    if not receipt_number:
        return jsonify({'success': False, 'error': '收件编号不能为空'}), 400
    
    if not applicants or len(applicants) == 0:
        return jsonify({'success': False, 'error': '至少需要一个申请人'}), 400
    
    conn = get_labor_db_connection()
    cursor = conn.cursor()
    
    try:
        # 1. 根据模式处理案件主表
        if mode == 'create':
            # 新建模式：检查收件编号是否已存在
            cursor.execute(
                "SELECT id FROM cases WHERE receipt_number = %s AND status = 1",
                (receipt_number,)
            )
            existing_case = cursor.fetchone()
            
            if existing_case:
                return jsonify({
                    'success': False, 
                    'error': f'收件编号 "{receipt_number}" 已存在，无法创建',
                    'code': 'DUPLICATE_RECEIPT_NUMBER'
                }), 409
            
            # 插入新案件
            cursor.execute(
                "INSERT INTO cases (receipt_number) VALUES (%s)",
                (receipt_number,)
            )
            case_id = cursor.lastrowid
            
        else:  # update 模式
            # 编辑模式：必须提供 case_id
            if not case_id:
                return jsonify({'success': False, 'error': '编辑模式必须提供案件ID'}), 400
            
            # 验证案件存在且状态正常
            cursor.execute(
                "SELECT id, receipt_number FROM cases WHERE id = %s AND status = 1",
                (case_id,)
            )
            existing_case = cursor.fetchone()
            
            if not existing_case:
                return jsonify({'success': False, 'error': '案件不存在或已被删除'}), 404
            
            # 如果修改了收件编号，检查新编号是否与其他案件冲突
            existing_receipt = existing_case[1]
            if receipt_number != existing_receipt:
                cursor.execute(
                    "SELECT id FROM cases WHERE receipt_number = %s AND status = 1 AND id != %s",
                    (receipt_number, case_id)
                )
                if cursor.fetchone():
                    return jsonify({
                        'success': False, 
                        'error': f'收件编号 "{receipt_number}" 已被其他案件使用',
                        'code': 'DUPLICATE_RECEIPT_NUMBER'
                    }), 409
            
            # 更新案件时间
            cursor.execute(
                "UPDATE cases SET receipt_number = %s, update_time = NOW() WHERE id = %s",
                (receipt_number, case_id)
            )
            
            # 删除旧的关联数据，重新插入
            cursor.execute("DELETE FROM arbitration_requests WHERE case_id = %s", (case_id,))
            cursor.execute("DELETE FROM evidence WHERE case_id = %s", (case_id,))
            cursor.execute("DELETE FROM respondents WHERE case_id = %s", (case_id,))
            cursor.execute("DELETE FROM applicants WHERE case_id = %s", (case_id,))
        
        # 2. 插入申请人及其仲裁请求
        # 建立 seq_no -> applicant_id 映射，用于证据关联
        applicant_seq_to_id = {}
        for applicant in applicants:
            cursor.execute("""
                INSERT INTO applicants (
                    case_id, seq_no, name, gender, nation, birth_date,
                    address, phone, id_card, employment_date, position, work_location,
                    monthly_salary, facts_reasons
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                case_id, applicant.get('seq_no'), applicant.get('name'),
                applicant.get('gender'), applicant.get('nation'), applicant.get('birth_date'),
                applicant.get('address'), applicant.get('phone'), applicant.get('id_card'),
                applicant.get('employment_date'), applicant.get('position'), applicant.get('work_location'),
                applicant.get('monthly_salary'), applicant.get('facts_reasons')
            ))
            applicant_id = cursor.lastrowid
            # 保存 seq_no 到 applicant_id 的映射
            applicant_seq_to_id[applicant.get('seq_no')] = applicant_id
            
            # 插入该申请人的仲裁请求
            for req in applicant.get('requests', []):
                cursor.execute("""
                    INSERT INTO arbitration_requests (applicant_id, case_id, seq_no, content)
                    VALUES (%s, %s, %s, %s)
                """, (applicant_id, case_id, req.get('seq_no'), req.get('content')))
        
        # 3. 插入被申请人
        for respondent in respondents:
            cursor.execute("""
                INSERT INTO respondents (
                    case_id, seq_no, name, legal_person, position,
                    address, phone, unified_code
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                case_id, respondent.get('seq_no'), respondent.get('name'),
                respondent.get('legal_person'), respondent.get('position'),
                respondent.get('address'), respondent.get('phone'), respondent.get('unified_code')
            ))
        
        # 4. 插入证据
        for evi in evidence:
            page_range = evi.get('page_range', '')
            page_start = ''
            page_end = ''
            if page_range and '-' in page_range:
                parts = page_range.split('-')
                page_start = parts[0]
                page_end = parts[1] if len(parts) > 1 else ''
            
            # 前端传来的是 applicant_seq_no，需要转换为真正的 applicant_id
            evi_applicant_seq = evi.get('applicant_seq_no')
            evi_applicant_id = applicant_seq_to_id.get(evi_applicant_seq) if evi_applicant_seq else None
            
            # 验证申请人存在性
            if evi_applicant_seq and evi_applicant_id is None:
                raise ValueError(f"证据 '{evi.get('name')}' 关联的申请人序号 {evi_applicant_seq} 不存在")
            
            cursor.execute("""
                INSERT INTO evidence (
                    case_id, applicant_id, seq_no, name, source, purpose, page_start, page_end, page_range
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                case_id, evi_applicant_id, evi.get('seq_no'), evi.get('name'), 
                evi.get('source'), evi.get('purpose'), page_start, page_end, page_range
            ))
        
        conn.commit()
        return jsonify({'success': True, 'case_id': case_id})
        
    except Exception as e:
        conn.rollback()
        logger.error(f"保存案件失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()


def get_case_detail_by_id(case_id):
    """根据案件ID获取完整详情（内部函数）"""
    conn = get_labor_db_connection()
    cursor = conn.cursor(pymysql.cursors.DictCursor)
    
    try:
        # 查询案件基本信息
        cursor.execute(
            "SELECT * FROM cases WHERE id = %s AND status = 1",
            (case_id,)
        )
        case = cursor.fetchone()
        
        if not case:
            return None
        
        # 查询申请人
        cursor.execute(
            "SELECT * FROM applicants WHERE case_id = %s ORDER BY seq_no",
            (case_id,)
        )
        applicants = cursor.fetchall()
        
        # 查询每个申请人的仲裁请求
        for applicant in applicants:
            cursor.execute(
                "SELECT seq_no, content FROM arbitration_requests WHERE applicant_id = %s ORDER BY seq_no",
                (applicant['id'],)
            )
            applicant['requests'] = cursor.fetchall()
        
        # 查询被申请人
        cursor.execute(
            "SELECT * FROM respondents WHERE case_id = %s ORDER BY seq_no",
            (case_id,)
        )
        respondents = cursor.fetchall()
        
        # 查询证据，并转换 applicant_id 为 applicant_seq_no
        cursor.execute(
            """SELECT e.*, a.seq_no as applicant_seq_no 
               FROM evidence e 
               LEFT JOIN applicants a ON e.applicant_id = a.id 
               WHERE e.case_id = %s 
               ORDER BY e.seq_no""",
            (case_id,)
        )
        evidence = cursor.fetchall()
        # 处理证据数据，将 applicant_id 替换为 applicant_seq_no
        for evi in evidence:
            evi['applicant_seq_no'] = evi.get('applicant_seq_no')  # 可能为None
        
        return {
            'case': case,
            'applicants': applicants,
            'respondents': respondents,
            'evidence': evidence
        }
        
    except Exception as e:
        logger.error(f"获取案件详情失败: {str(e)}")
        return None
    finally:
        cursor.close()
        conn.close()


@app.route('/api/cases/query', methods=['GET'])
def query_case():
    """根据收件编号查询案件"""
    receipt_number = request.args.get('receipt_number')
    
    if not receipt_number:
        return jsonify({'success': False, 'error': '收件编号不能为空'}), 400
    
    conn = get_labor_db_connection()
    cursor = conn.cursor(pymysql.cursors.DictCursor)
    
    try:
        # 查询案件完整信息（关联查询）
        cursor.execute(
            "SELECT * FROM cases WHERE receipt_number = %s AND status = 1",
            (receipt_number,)
        )
        case = cursor.fetchone()
        
        if not case:
            return jsonify({'success': False, 'error': '案件不存在'}), 404
        
        case_id = case['id']
        
        # 查询申请人
        cursor.execute(
            "SELECT * FROM applicants WHERE case_id = %s ORDER BY seq_no",
            (case_id,)
        )
        applicants = cursor.fetchall()
        
        # 查询每个申请人的仲裁请求
        for applicant in applicants:
            cursor.execute(
                "SELECT seq_no, content FROM arbitration_requests WHERE applicant_id = %s ORDER BY seq_no",
                (applicant['id'],)
            )
            applicant['requests'] = cursor.fetchall()
        
        # 查询被申请人
        cursor.execute(
            "SELECT * FROM respondents WHERE case_id = %s ORDER BY seq_no",
            (case_id,)
        )
        respondents = cursor.fetchall()
        
        # 查询证据
        cursor.execute(
            "SELECT * FROM evidence WHERE case_id = %s ORDER BY seq_no",
            (case_id,)
        )
        evidence = cursor.fetchall()
        
        return jsonify({
            'success': True,
            'data': {
                'case': case,
                'applicants': applicants,
                'respondents': respondents,
                'evidence': evidence
            }
        })
        
    except Exception as e:
        logger.error(f"查询案件失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


@app.route('/api/cases/<int:case_id>', methods=['GET'])
def get_case_by_id(case_id):
    """根据案件ID查询案件详情（用于编辑）"""
    data = get_case_detail_by_id(case_id)
    
    if not data:
        return jsonify({'success': False, 'error': '案件不存在或已被删除'}), 404
    
    return jsonify({'success': True, 'data': data})


@app.route('/api/cases/list', methods=['GET'])
def list_cases():
    """获取案件列表"""
    page = int(request.args.get('page', 1))
    page_size = int(request.args.get('page_size', 10))
    offset = (page - 1) * page_size
    
    conn = get_labor_db_connection()
    cursor = conn.cursor(pymysql.cursors.DictCursor)
    
    try:
        # 查询案件列表（带统计信息）
        cursor.execute("""
            SELECT 
                c.id, c.receipt_number, c.create_time, c.update_time,
                COUNT(DISTINCT a.id) AS applicant_count,
                COUNT(DISTINCT r.id) AS respondent_count,
                COUNT(DISTINCT e.id) AS evidence_count,
                GROUP_CONCAT(DISTINCT a.name ORDER BY a.seq_no) AS applicant_names
            FROM cases c
            LEFT JOIN applicants a ON c.id = a.case_id
            LEFT JOIN respondents r ON c.id = r.case_id
            LEFT JOIN evidence e ON c.id = e.case_id
            WHERE c.status = 1
            GROUP BY c.id
            ORDER BY c.create_time DESC
            LIMIT %s OFFSET %s
        """, (page_size, offset))
        cases = cursor.fetchall()
        
        # 查询总数
        cursor.execute("SELECT COUNT(*) as total FROM cases WHERE status = 1")
        total = cursor.fetchone()['total']
        
        return jsonify({
            'success': True,
            'data': {
                'list': cases,
                'total': total,
                'page': page,
                'page_size': page_size
            }
        })
        
    except Exception as e:
        logger.error(f"获取案件列表失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()


@app.route('/api/cases/<int:case_id>', methods=['DELETE'])
def delete_case(case_id):
    """删除案件（软删除）"""
    conn = get_labor_db_connection()
    cursor = conn.cursor()
    
    try:
        # 先检查案件是否存在且未删除
        cursor.execute(
            "SELECT id FROM cases WHERE id = %s AND status = 1",
            (case_id,)
        )
        case = cursor.fetchone()
        
        if not case:
            return jsonify({'success': False, 'error': '案件不存在或已被删除'}), 404
        
        # 执行软删除
        cursor.execute(
            "UPDATE cases SET status = 0 WHERE id = %s",
            (case_id,)
        )
        conn.commit()
        return jsonify({'success': True, 'message': '案件已删除'})
    except Exception as e:
        conn.rollback()
        logger.error(f"删除案件失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()


@app.route('/api/cases/<int:case_id>/applicants', methods=['GET'])
def get_applicants(case_id):
    """获取案件的所有申请人"""
    conn = get_labor_db_connection()
    cursor = conn.cursor(pymysql.cursors.DictCursor)
    
    try:
        cursor.execute(
            "SELECT * FROM applicants WHERE case_id = %s ORDER BY seq_no",
            (case_id,)
        )
        applicants = cursor.fetchall()
        return jsonify({'success': True, 'data': applicants})
    except Exception as e:
        logger.error(f"获取申请人列表失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()


@app.route('/api/applicants/<int:applicant_id>', methods=['GET'])
def get_applicant_detail(applicant_id):
    """获取单个申请人详情（包含仲裁请求）"""
    conn = get_labor_db_connection()
    cursor = conn.cursor(pymysql.cursors.DictCursor)
    
    try:
        cursor.execute(
            "SELECT * FROM applicants WHERE id = %s",
            (applicant_id,)
        )
        applicant = cursor.fetchone()
        
        if not applicant:
            return jsonify({'success': False, 'error': '申请人不存在'}), 404
        
        cursor.execute(
            "SELECT seq_no, content FROM arbitration_requests WHERE applicant_id = %s ORDER BY seq_no",
            (applicant_id,)
        )
        applicant['requests'] = cursor.fetchall()
        
        return jsonify({'success': True, 'data': applicant})
    except Exception as e:
        logger.error(f"获取申请人详情失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()


@app.route('/api/applicants/<int:applicant_id>/requests', methods=['GET'])
def get_applicant_requests(applicant_id):
    """获取申请人的仲裁请求"""
    conn = get_labor_db_connection()
    cursor = conn.cursor(pymysql.cursors.DictCursor)
    
    try:
        cursor.execute(
            "SELECT * FROM arbitration_requests WHERE applicant_id = %s ORDER BY seq_no",
            (applicant_id,)
        )
        requests = cursor.fetchall()
        return jsonify({'success': True, 'data': requests})
    except Exception as e:
        logger.error(f"获取仲裁请求失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()


@app.route('/api/cases/<int:case_id>/respondents', methods=['GET'])
def get_respondents(case_id):
    """获取案件的所有被申请人"""
    conn = get_labor_db_connection()
    cursor = conn.cursor(pymysql.cursors.DictCursor)
    
    try:
        cursor.execute(
            "SELECT * FROM respondents WHERE case_id = %s ORDER BY seq_no",
            (case_id,)
        )
        respondents = cursor.fetchall()
        return jsonify({'success': True, 'data': respondents})
    except Exception as e:
        logger.error(f"获取被申请人列表失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()


@app.route('/api/cases/<int:case_id>/evidence', methods=['GET'])
def get_evidence(case_id):
    """获取案件的所有证据"""
    conn = get_labor_db_connection()
    cursor = conn.cursor(pymysql.cursors.DictCursor)
    
    try:
        cursor.execute(
            "SELECT * FROM evidence WHERE case_id = %s ORDER BY seq_no",
            (case_id,)
        )
        evidence = cursor.fetchall()
        return jsonify({'success': True, 'data': evidence})
    except Exception as e:
        logger.error(f"获取证据列表失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        cursor.close()
        conn.close()


# ============================================
# 错误处理
# ============================================

@app.errorhandler(404)
def not_found(error):
    return jsonify({
        'code': 404,
        'message': 'API端点不存在',
        'timestamp': datetime.now().isoformat()
    }), 404


@app.errorhandler(405)
def method_not_allowed(error):
    return jsonify({
        'code': 405,
        'message': '请求方法不允许',
        'timestamp': datetime.now().isoformat()
    }), 405


@app.errorhandler(500)
def internal_server_error(error):
    logger.error(f"服务器内部错误: {str(error)}")
    return jsonify({
        'code': 500,
        'message': '服务器内部错误',
        'timestamp': datetime.now().isoformat()
    }), 500


# ============================================
# 收件查询API（代理内部服务）
# ============================================

INTERNAL_API_BASE = "http://10.96.10.78:8080/v1/api/admin/arb/receive"

@app.route('/api/receive/query', methods=['GET'])
def query_receive():
    """
    收件查询接口 - 代理调用内部服务
    支持参数：
    - page: 页码（默认1）
    - page_size: 每页数量（默认20）
    - application_date: 申请日期（格式：YYYY-MM-DD）
    - status: 状态（0-收件登记, 1-审核通过, 2-审核不通过, 3-审批通过, 4-审批不通过, 5-提交）
    - case_no: 案件编号（用户输入202691，实际传参[2026]91）
    - search: 关键字（搜索申请人/被申请人/案由）
    """
    try:
        # 获取查询参数
        page = int(request.args.get('page', 1))
        page_size = int(request.args.get('page_size', 20))
        application_date = request.args.get('application_date')
        status = request.args.get('status')
        case_no = request.args.get('case_no')
        search = request.args.get('search')
        
        # 构建内部API请求参数（内部API使用从1开始的页码）
        params = {
            'page': page,  # 内部API使用从1开始的页码
            'page_size': page_size
        }
        
        if application_date:
            params['application_date'] = application_date
            
        if status:
            params['status'] = status
            logger.info(f"状态筛选: {status}")
        
        if search:
            params['search'] = search.strip()
            logger.info(f"关键字筛选: {search}")
            
        if case_no:
            # 支持多种输入格式，统一转换为 [2026]91 格式
            original_case_no = case_no.strip()
            converted_case_no = original_case_no
            
            # 提取方括号中的内容
            import re
            match = re.search(r'\[(\d{4})\](\d+)', original_case_no)
            if match:
                # 已经是 [2026]91 或包含这种格式
                converted_case_no = f"[{match.group(1)}]{match.group(2)}"
            elif original_case_no.isdigit() and len(original_case_no) >= 6:
                # 纯数字格式 202691
                year = original_case_no[:4]
                num = original_case_no[4:]
                converted_case_no = f"[{year}]{num}"
            elif re.match(r'^\d{4}\]\d+$', original_case_no):
                # 2026]91 格式，补充左括号
                converted_case_no = f"[{original_case_no}"
            
            params['case_no'] = converted_case_no
            logger.info(f"案件编号转换: {original_case_no} -> {converted_case_no}")
        
        # 检查并更新登录状态
        if not login_manager.check_and_renew_login():
            logger.error("获取有效登录信息失败")
            return jsonify({
                'code': 401,
                'message': '登录失败，无法查询收件信息',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 401
        
        # 获取带认证信息的请求头
        headers = login_manager.get_auth_headers()
        if not headers:
            logger.error("获取认证请求头失败")
            return jsonify({
                'code': 401,
                'message': '获取认证信息失败',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 401
        
        logger.info(f"收件查询请求: {INTERNAL_API_BASE}")
        logger.info(f"  原始参数: page={page}, page_size={page_size}, application_date={application_date}, status={status}, case_no={case_no}")
        logger.info(f"  转换后参数: {params}")
        
        # 调用内部服务（带认证头）
        response = requests.get(
            INTERNAL_API_BASE,
            headers=headers,
            params=params,
            timeout=30
        )
        
        # 如果返回401，尝试重新登录后重试
        if response.status_code == 401:
            logger.warning("收件查询: 内部服务返回401，尝试重新登录...")
            login_manager.current_auth_key = None
            login_manager.current_session_id = None
            
            if not login_manager.check_and_renew_login():
                logger.error("收件查询: 重新登录失败")
                return jsonify({
                    'code': 401,
                    'message': '登录已失效，请刷新页面重试',
                    'data': None,
                    'timestamp': datetime.now().isoformat()
                }), 401
            
            # 获取新的认证头并重试
            headers = login_manager.get_auth_headers()
            response = requests.get(
                INTERNAL_API_BASE,
                headers=headers,
                params=params,
                timeout=30
            )
        
        # 返回结果
        if response.status_code == 200:
            data = response.json()
            
            # 处理数据：将案由平铺到 item 级别 + 内存分页
            # 数据结构: {code:200, data:{code:200, data:{data:[...], totalNum:...}}}
            try:
                outer_data = data.get('data', {})
                if isinstance(outer_data, dict):
                    inner_data = outer_data.get('data', {})
                    if isinstance(inner_data, dict):
                        items = inner_data.get('data', [])
                        if isinstance(items, list):
                            # 保存总条数
                            total_count = len(items)
                            # 内存分页
                            offset = (page - 1) * page_size
                            paginated_items = items[offset:offset + page_size]
                            # 更新数据
                            inner_data['data'] = paginated_items
                            inner_data['totalNum'] = total_count
                            # 更新案由
                            logger.info(f"处理案由: {len(paginated_items)}条数据")
                            for item in paginated_items:
                                if isinstance(item, dict):
                                    cases = item.get('cases', {})
                                    if cases and isinstance(cases, dict):
                                        case_reason = cases.get('case_reason')
                                        if case_reason:
                                            logger.info(f"  找到案由: {case_reason[:20]}")
                                            if not item.get('case_reason'):
                                                item['case_reason'] = case_reason
                                                logger.info(f"  已设置案由")
                                        else:
                                            logger.info(f"  无案由")
            except Exception as e:
                logger.warning(f"处理数据时出错: {e}")
            
            return jsonify({
                'code': 200,
                'message': '查询成功',
                'data': data,
                'timestamp': datetime.now().isoformat()
            })
        else:
            logger.error(f"内部服务返回错误: {response.status_code}, {response.text}")
            return jsonify({
                'code': response.status_code,
                'message': f'内部服务错误: {response.status_code}',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 500
            
    except requests.exceptions.Timeout:
        logger.error("收件查询超时")
        return jsonify({
            'code': 504,
            'message': '请求内部服务超时',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 504
    except requests.exceptions.RequestException as e:
        logger.error(f"收件查询请求异常: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'请求内部服务失败: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500
    except Exception as e:
        logger.error(f"收件查询错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器错误: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


@app.route('/api/receive/detail', methods=['GET'])
def query_receive_detail():
    """
    收件详情查询接口 - 代理调用内部服务
    参数:
      - id: 收件记录ID（必填）
    """
    try:
        # 获取查询参数
        item_id = request.args.get('id')
        
        if not item_id:
            return jsonify({
                'code': 400,
                'message': '缺少参数: id',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 400
        
        # 检查并更新登录状态
        if not login_manager.check_and_renew_login():
            logger.error("获取有效登录信息失败")
            return jsonify({
                'code': 401,
                'message': '登录失败，无法查询详情',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 401
        
        # 获取带认证信息的请求头
        headers = login_manager.get_auth_headers()
        if not headers:
            logger.error("获取认证请求头失败")
            return jsonify({
                'code': 401,
                'message': '获取认证信息失败',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 401
        
        # 构建内部API URL
        detail_url = f"{INTERNAL_API_BASE.replace('/receive', '')}/{item_id}/receive"
        
        logger.info(f"收件详情查询请求: {detail_url}")
        
        # 调用内部服务（带认证头）
        response = requests.get(
            detail_url,
            headers=headers,
            timeout=30
        )
        
        # 如果返回401，尝试重新登录后重试
        if response.status_code == 401:
            logger.warning("收件详情: 内部服务返回401，尝试重新登录...")
            login_manager.current_auth_key = None
            login_manager.current_session_id = None
            
            if not login_manager.check_and_renew_login():
                logger.error("收件详情: 重新登录失败")
                return jsonify({
                    'code': 401,
                    'message': '登录已失效，请刷新页面重试',
                    'data': None,
                    'timestamp': datetime.now().isoformat()
                }), 401
            
            # 获取新的认证头并重试
            headers = login_manager.get_auth_headers()
            response = requests.get(
                detail_url,
                headers=headers,
                timeout=30
            )
        
        # 返回结果
        if response.status_code == 200:
            data = response.json()
            
            # 处理数据：将案由平铺到 item 级别，方便前端显示
            if data and isinstance(data, dict) and 'data' in data:
                items = data['data']
                if isinstance(items, list):
                    for item in items:
                        if isinstance(item, dict):
                            # 从 cases 中提取案由到 item 级别
                            cases = item.get('cases', {})
                            if cases and isinstance(cases, dict):
                                if not item.get('case_reason') and cases.get('case_reason'):
                                    item['case_reason'] = cases.get('case_reason')
            
            return jsonify({
                'code': 200,
                'message': '查询成功',
                'data': data,
                'timestamp': datetime.now().isoformat()
            })
        else:
            logger.error(f"内部服务返回错误: {response.status_code}, {response.text}")
            return jsonify({
                'code': response.status_code,
                'message': f'内部服务错误: {response.status_code}',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 500
            
    except requests.exceptions.Timeout:
        logger.error("收件详情查询超时")
        return jsonify({
            'code': 504,
            'message': '请求内部服务超时',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 504
    except requests.exceptions.RequestException as e:
        logger.error(f"收件详情查询请求异常: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'请求内部服务失败: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500
    except Exception as e:
        logger.error(f"收件详情查询错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器错误: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


# ============================================
# 预约仲裁申请查询API（代理内部服务）
# ============================================

RESERVE_API_BASE = "http://10.96.10.78:8080/v1/api/admin/arb/reserve"

@app.route('/api/reserve/query', methods=['GET'])
def query_reserve():
    """
    预约仲裁申请查询接口 - 代理调用内部服务
    支持参数：
    - type: 类型（默认1）
    - page: 页码（默认1）
    - page_size: 每页数量（默认10）
    - status: 状态（0-待审核, 1-审核通过, 2-审核不通过, 3-现场确认同意, 4-现场确认不同意）
    - applicant: 申请人姓名
    - respondent: 被申请人名称
    - submit_at: 申请日期（格式：YYYY-MM-DD）
    """
    try:
        # 获取查询参数
        type_val = request.args.get('type', '1')
        page = int(request.args.get('page', 1))
        page_size = int(request.args.get('page_size', 10))
        status = request.args.get('status')
        applicant = request.args.get('applicant')
        respondent = request.args.get('respondent')
        submit_at = request.args.get('submit_at')
        
        # 构建内部API请求参数（内部API使用从1开始的页码）
        params = {
            'type': type_val,
            'page': page,           # 内部API使用从1开始的页码
            'page_size': page_size
        }
        
        if status:
            params['status'] = status
            logger.info(f"状态筛选: {status}")
        
        if applicant:
            params['applicant'] = applicant.strip()
            logger.info(f"申请人筛选: {applicant}")
            
        if respondent:
            params['respondent'] = respondent.strip()
            logger.info(f"被申请人筛选: {respondent}")
            
        if submit_at:
            params['submit_at'] = submit_at
            logger.info(f"申请日期筛选: {submit_at}")
        
        # 检查并更新登录状态
        if not login_manager.check_and_renew_login():
            logger.error("获取有效登录信息失败")
            return jsonify({
                'code': 401,
                'message': '登录失败，无法查询预约仲裁申请信息',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 401
        
        # 获取带认证信息的请求头
        headers = login_manager.get_auth_headers()
        if not headers:
            logger.error("获取认证请求头失败")
            return jsonify({
                'code': 401,
                'message': '获取认证信息失败',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 401
        
        logger.info(f"预约仲裁申请查询请求: {RESERVE_API_BASE}")
        logger.info(f"  参数: {params}")
        
        # 调用内部服务（带认证头）
        response = requests.get(
            RESERVE_API_BASE,
            headers=headers,
            params=params,
            timeout=30
        )
        
        # 如果返回401，尝试重新登录后重试
        if response.status_code == 401:
            logger.warning("预约查询: 内部服务返回401，尝试重新登录...")
            login_manager.current_auth_key = None
            login_manager.current_session_id = None
            
            if not login_manager.check_and_renew_login():
                logger.error("预约查询: 重新登录失败")
                return jsonify({
                    'code': 401,
                    'message': '登录已失效，请刷新页面重试',
                    'data': None,
                    'timestamp': datetime.now().isoformat()
                }), 401
            
            # 获取新的认证头并重试
            headers = login_manager.get_auth_headers()
            response = requests.get(
                RESERVE_API_BASE,
                headers=headers,
                params=params,
                timeout=30
            )
        
        # 返回结果
        if response.status_code == 200:
            data = response.json()
            
            # 调试：打印返回的数据ID列表
            try:
                # 数据结构: {code:200, data:{code:200, data:{data:[...], totalNum:...}}}
                outer_data = data.get('data', {})
                if isinstance(outer_data, dict):
                    inner_data = outer_data.get('data', {})
                    if isinstance(inner_data, dict):
                        items = inner_data.get('data', [])
                        total_num = inner_data.get('totalNum', 0)
                        if isinstance(items, list) and len(items) > 0:
                            ids = [item.get('id') for item in items[0:5]]
                            case_nos = [item.get('case_no') for item in items[0:5]]
                            logger.info(f"立案查询返回: page={page}, 数据条数={len(items)}, 总条数={total_num}, "
                                       f"前5条ID={ids}, 前5条案号={case_nos}")
            except Exception as e:
                logger.warning(f"调试日志出错: {e}")
            
            # 处理数据：将案由平铺到 item 级别
            # 数据结构: {code:200, data:{code:200, data:{data:[...], totalNum:...}}}
            try:
                outer_data = data.get('data', {})
                if isinstance(outer_data, dict):
                    inner_data = outer_data.get('data', {})
                    if isinstance(inner_data, dict):
                        items = inner_data.get('data', [])
                        if isinstance(items, list):
                            for item in items:
                                if isinstance(item, dict):
                                    cases = item.get('cases', {})
                                    if cases and isinstance(cases, dict):
                                        case_reason = cases.get('case_reason')
                                        if case_reason and not item.get('case_reason'):
                                            item['case_reason'] = case_reason
            except Exception as e:
                logger.warning(f"处理案由数据时出错: {e}")
            
            return jsonify({
                'code': 200,
                'message': '查询成功',
                'data': data,
                'timestamp': datetime.now().isoformat()
            })
        else:
            logger.error(f"内部服务返回错误: {response.status_code}, {response.text}")
            return jsonify({
                'code': response.status_code,
                'message': f'内部服务错误: {response.status_code}',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 500
            
    except requests.exceptions.Timeout:
        logger.error("预约仲裁申请查询超时")
        return jsonify({
            'code': 504,
            'message': '请求内部服务超时',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 504
    except requests.exceptions.RequestException as e:
        logger.error(f"预约仲裁申请查询请求异常: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'请求内部服务失败: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500
    except Exception as e:
        logger.error(f"预约仲裁申请查询错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器错误: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


@app.route('/api/reserve/detail', methods=['GET'])
def query_reserve_detail():
    """
    预约仲裁申请详情查询接口 - 代理调用内部服务
    参数:
      - id: 预约申请记录ID（必填）
    """
    try:
        # 获取查询参数
        item_id = request.args.get('id')
        
        if not item_id:
            return jsonify({
                'code': 400,
                'message': '缺少参数: id',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 400
        
        # 检查并更新登录状态
        if not login_manager.check_and_renew_login():
            logger.error("获取有效登录信息失败")
            return jsonify({
                'code': 401,
                'message': '登录失败，无法查询详情',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 401
        
        # 获取带认证信息的请求头
        headers = login_manager.get_auth_headers()
        if not headers:
            logger.error("获取认证请求头失败")
            return jsonify({
                'code': 401,
                'message': '获取认证信息失败',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 401
        
        # 构建内部API URL: /v1/api/admin/arb/{id}/reserve
        detail_url = f"{RESERVE_API_BASE.replace('/reserve', '')}/{item_id}/reserve"
        
        logger.info(f"预约仲裁申请详情查询请求: {detail_url}")
        
        # 调用内部服务（带认证头）
        response = requests.get(
            detail_url,
            headers=headers,
            timeout=30
        )
        
        # 如果返回401，尝试重新登录后重试
        if response.status_code == 401:
            logger.warning("预约详情: 内部服务返回401，尝试重新登录...")
            login_manager.current_auth_key = None
            login_manager.current_session_id = None
            
            if not login_manager.check_and_renew_login():
                logger.error("预约详情: 重新登录失败")
                return jsonify({
                    'code': 401,
                    'message': '登录已失效，请刷新页面重试',
                    'data': None,
                    'timestamp': datetime.now().isoformat()
                }), 401
            
            # 获取新的认证头并重试
            headers = login_manager.get_auth_headers()
            response = requests.get(
                detail_url,
                headers=headers,
                timeout=30
            )
        
        # 返回结果
        if response.status_code == 200:
            data = response.json()
            return jsonify({
                'code': 200,
                'message': '查询成功',
                'data': data,
                'timestamp': datetime.now().isoformat()
            })
        else:
            logger.error(f"内部服务返回错误: {response.status_code}, {response.text}")
            return jsonify({
                'code': response.status_code,
                'message': f'内部服务错误: {response.status_code}',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 500
            
    except requests.exceptions.Timeout:
        logger.error("预约仲裁申请详情查询超时")
        return jsonify({
            'code': 504,
            'message': '请求内部服务超时',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 504
    except requests.exceptions.RequestException as e:
        logger.error(f"预约仲裁申请详情查询请求异常: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'请求内部服务失败: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500
    except Exception as e:
        logger.error(f"预约仲裁申请详情查询错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器错误: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


# ============================================
# 立案查询API（代理内部服务）
# ============================================

HANDLE_API_BASE = "http://10.96.10.78:8080/v1/api/admin/arb/handle"

@app.route('/api/handle/query', methods=['GET'])
def query_handle():
    """
    立案查询接口 - 代理调用内部服务
    支持参数：
    - page: 页码（默认1）
    - page_size: 每页数量（默认20）
    - case_no: 案件编号（用户输入202691，实际传参[2026]91）
    - search: 关键字（搜索申请人/被申请人/案由）
    """
    try:
        # 获取查询参数
        page = int(request.args.get('page', 1))
        page_size = int(request.args.get('page_size', 20))
        case_no = request.args.get('case_no')
        search = request.args.get('search')
        
        # 构建内部API请求参数（内部API使用从1开始的页码）
        params = {
            'page': page,  # 内部API使用从1开始的页码
            'page_size': page_size
        }
        
        if search:
            params['search'] = search.strip()
            logger.info(f"立案查询关键字筛选: {search}")
            
        if case_no:
            # 支持多种输入格式，统一转换为 [2026]91 格式
            original_case_no = case_no.strip()
            converted_case_no = original_case_no
            
            # 提取方括号中的内容
            import re
            match = re.search(r'\[(\d{4})\](\d+)', original_case_no)
            if match:
                # 已经是 [2026]91 或包含这种格式
                converted_case_no = f"[{match.group(1)}]{match.group(2)}"
            elif original_case_no.isdigit() and len(original_case_no) >= 6:
                # 纯数字格式 202691
                year = original_case_no[:4]
                num = original_case_no[4:]
                converted_case_no = f"[{year}]{num}"
            elif re.match(r'^\d{4}\]\d+$', original_case_no):
                # 2026]91 格式，补充左括号
                converted_case_no = f"[{original_case_no}"
            
            params['case_no'] = converted_case_no
            logger.info(f"立案查询案件编号转换: {original_case_no} -> {converted_case_no}")
        
        # 检查并更新登录状态
        if not login_manager.check_and_renew_login():
            logger.error("获取有效登录信息失败")
            return jsonify({
                'code': 401,
                'message': '登录失败，无法查询立案信息',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 401
        
        # 获取带认证信息的请求头
        headers = login_manager.get_auth_headers()
        if not headers:
            logger.error("获取认证请求头失败")
            return jsonify({
                'code': 401,
                'message': '获取认证信息失败',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 401
        
        logger.info(f"立案查询请求: {HANDLE_API_BASE}")
        logger.info(f"  原始参数: page={page}, page_size={page_size}, case_no={case_no}")
        logger.info(f"  转换后参数: {params}")
        
        # 调用内部服务（带认证头）
        response = requests.get(
            HANDLE_API_BASE,
            headers=headers,
            params=params,
            timeout=30
        )
        
        # 如果返回401，尝试重新登录后重试
        if response.status_code == 401:
            logger.warning("内部服务返回401，尝试重新登录...")
            login_manager.current_auth_key = None
            login_manager.current_session_id = None
            
            if not login_manager.check_and_renew_login():
                logger.error("重新登录失败")
                return jsonify({
                    'code': 401,
                    'message': '登录已失效，请刷新页面重试',
                    'data': None,
                    'timestamp': datetime.now().isoformat()
                }), 401
            
            # 获取新的认证头并重试
            headers = login_manager.get_auth_headers()
            response = requests.get(
                HANDLE_API_BASE,
                headers=headers,
                params=params,
                timeout=30
            )
        
        # 返回结果
        if response.status_code == 200:
            data = response.json()
            
            # 处理数据：将案由平铺到 item 级别 + 内存分页
            # 数据结构: {code:200, data:{code:200, data:{data:[...], totalNum:...}}}
            try:
                outer_data = data.get('data', {})
                if isinstance(outer_data, dict):
                    inner_data = outer_data.get('data', {})
                    if isinstance(inner_data, dict):
                        items = inner_data.get('data', [])
                        if isinstance(items, list):
                            # 保存总条数
                            total_count = len(items)
                            # 内存分页
                            offset = (page - 1) * page_size
                            paginated_items = items[offset:offset + page_size]
                            # 更新数据
                            inner_data['data'] = paginated_items
                            inner_data['totalNum'] = total_count
                            # 更新案由
                            print(f"[DEBUG] 处理案由: {len(paginated_items)}条", flush=True)
                            for item in paginated_items:
                                if isinstance(item, dict):
                                    cases = item.get('cases')
                                    case_reason = None
                                    if cases:
                                        if isinstance(cases, dict):
                                            # cases 是对象
                                            case_reason = cases.get('case_reason')
                                        elif isinstance(cases, list) and len(cases) > 0:
                                            # cases 是数组，取第一个元素
                                            first_case = cases[0]
                                            if isinstance(first_case, dict):
                                                case_reason = first_case.get('case_reason')
                                    if case_reason:
                                        print(f"[DEBUG] 设置案由: {case_reason[:20]}", flush=True)
                                        item['case_reason'] = case_reason
            except Exception as e:
                logger.warning(f"处理数据时出错: {e}")
            
            return jsonify({
                'code': 200,
                'message': '查询成功',
                'data': data,
                'timestamp': datetime.now().isoformat()
            })
        else:
            logger.error(f"内部服务返回错误: {response.status_code}, {response.text}")
            return jsonify({
                'code': response.status_code,
                'message': f'内部服务错误: {response.status_code}',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 500
            
    except requests.exceptions.Timeout:
        logger.error("立案查询超时")
        return jsonify({
            'code': 504,
            'message': '请求内部服务超时',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 504
    except requests.exceptions.RequestException as e:
        logger.error(f"立案查询请求异常: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'请求内部服务失败: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500
    except Exception as e:
        logger.error(f"立案查询错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器错误: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


@app.route('/api/handle/detail', methods=['GET'])
def query_handle_detail():
    """
    立案详情查询接口 - 代理调用内部服务
    参数:
      - id: 案件ID（必填）
    """
    try:
        # 获取查询参数
        item_id = request.args.get('id')
        
        if not item_id:
            return jsonify({
                'code': 400,
                'message': '缺少参数: id',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 400
        
        # 检查并更新登录状态
        if not login_manager.check_and_renew_login():
            logger.error("获取有效登录信息失败")
            return jsonify({
                'code': 401,
                'message': '登录失败，无法查询详情',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 401
        
        # 获取带认证信息的请求头
        headers = login_manager.get_auth_headers()
        if not headers:
            logger.error("获取认证请求头失败")
            return jsonify({
                'code': 401,
                'message': '获取认证信息失败',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 401
        
        # 构建内部API URL
        detail_url = "http://10.96.10.78:8080/v1/api/admin/case/caseData"
        # 获取 case_material 的接口
        material_url = f"http://10.96.10.78:8080/v1/api/admin/arb/{item_id}/handle"
        
        logger.info(f"立案详情查询请求: {detail_url}")
        logger.info(f"  案件ID: {item_id}")
        
        # 并行调用两个接口
        from concurrent.futures import ThreadPoolExecutor, as_completed
        
        def fetch_case_data():
            """获取案件详情数据"""
            resp = requests.post(
                detail_url,
                headers=headers,
                json={'id': item_id},
                timeout=30
            )
            return ('case_data', resp)
        
        def fetch_material_data():
            """获取 case_material 数据"""
            resp = requests.get(
                material_url,
                headers=headers,
                timeout=30
            )
            return ('material', resp)
        
        # 执行两个请求
        case_response = None
        material_response = None
        
        with ThreadPoolExecutor(max_workers=2) as executor:
            futures = [
                executor.submit(fetch_case_data),
                executor.submit(fetch_material_data)
            ]
            for future in as_completed(futures):
                try:
                    data_type, resp = future.result()
                    if data_type == 'case_data':
                        case_response = resp
                    else:
                        material_response = resp
                except Exception as e:
                    logger.error(f"并行请求异常: {e}")
        
        # 使用案件详情响应作为主响应
        response = case_response
        
        # 如果返回401，尝试重新登录后重试
        if response.status_code == 401:
            logger.warning("内部服务返回401，尝试重新登录...")
            login_manager.current_auth_key = None
            login_manager.current_session_id = None
            
            if not login_manager.check_and_renew_login():
                logger.error("重新登录失败")
                return jsonify({
                    'code': 401,
                    'message': '登录已失效，请刷新页面重试',
                    'data': None,
                    'timestamp': datetime.now().isoformat()
                }), 401
            
            # 获取新的认证头并重试
            headers = login_manager.get_auth_headers()
            response = requests.post(
                detail_url,
                headers=headers,
                json={'id': item_id},
                timeout=30
            )
        
        # 返回结果
        if response.status_code == 200:
            data = response.json()
            
            # 提取并合并 case_material 数据
            case_material = None
            if material_response and material_response.status_code == 200:
                try:
                    material_data = material_response.json()
                    # 从 arb/{id}/handle 接口返回的数据中提取 case_material
                    if isinstance(material_data, dict):
                        # 尝试从不同层级获取 case_material
                        if 'case_material' in material_data:
                            case_material = material_data['case_material']
                        elif 'data' in material_data and isinstance(material_data['data'], dict):
                            # 检查 data 层级
                            if 'case_material' in material_data['data']:
                                case_material = material_data['data']['case_material']
                            # 检查 data.data 层级 (arb/{id}/handle 接口的实际结构)
                            elif 'data' in material_data['data'] and isinstance(material_data['data']['data'], dict):
                                case_material = material_data['data']['data'].get('case_material')
                        
                        if case_material is not None:
                            logger.info(f"成功获取 case_material，数据条数: {len(case_material) if isinstance(case_material, list) else 'N/A'}")
                except Exception as e:
                    logger.error(f"解析 case_material 数据失败: {e}")
            else:
                logger.warning(f"获取 case_material 失败，状态码: {material_response.status_code if material_response else 'None'}")
            
            # 将 case_material 合并到返回数据中
            if isinstance(data, dict):
                if 'data' in data and isinstance(data['data'], dict):
                    data['data']['case_material'] = case_material
                else:
                    data['case_material'] = case_material
            
            return jsonify({
                'code': 200,
                'message': '查询成功',
                'data': data,
                'timestamp': datetime.now().isoformat()
            })
        else:
            logger.error(f"内部服务返回错误: {response.status_code}, {response.text}")
            return jsonify({
                'code': response.status_code,
                'message': f'内部服务错误: {response.status_code}',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 500
            
    except requests.exceptions.Timeout:
        logger.error("立案详情查询超时")
        return jsonify({
            'code': 504,
            'message': '请求内部服务超时',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 504
    except requests.exceptions.RequestException as e:
        logger.error(f"立案详情查询请求异常: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'请求内部服务失败: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500
    except Exception as e:
        logger.error(f"立案详情查询错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器错误: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


# ============================================
# 健康检查
# ============================================

@app.route('/api/health', methods=['GET'])
def health_check():
    """健康检查接口"""
    return jsonify({'status': 'ok', 'timestamp': datetime.now().isoformat()})


@app.route('/api/file/proxy', methods=['GET'])
def proxy_file():
    """
    文件代理下载接口 - 用于代理访问内部服务器的文件
    解决客户端无法直接访问内网文件的问题
    参数:
      - url: 文件URL（必填，需要进行URL编码）
    """
    try:
        # 获取文件URL参数
        file_url = request.args.get('url')
        
        if not file_url:
            return jsonify({
                'code': 400,
                'message': '缺少参数: url',
                'timestamp': datetime.now().isoformat()
            }), 400
        
        # URL解码
        from urllib.parse import unquote
        file_url = unquote(file_url)
        
        # 验证URL是否来自允许的服务器（安全校验）
        allowed_hosts = ['10.96.10.78:8080', '10.96.10.78']
        is_allowed = any(host in file_url for host in allowed_hosts)
        
        if not is_allowed:
            logger.warning(f"尝试访问不允许的文件地址: {file_url}")
            return jsonify({
                'code': 403,
                'message': '访问被拒绝: 不允许的文件地址',
                'timestamp': datetime.now().isoformat()
            }), 403
        
        # 检查并更新登录状态（如果需要认证）
        if not login_manager.check_and_renew_login():
            logger.error("获取有效登录信息失败")
            return jsonify({
                'code': 401,
                'message': '登录失败，无法下载文件',
                'timestamp': datetime.now().isoformat()
            }), 401
        
        # 获取带认证信息的请求头
        headers = login_manager.get_auth_headers()
        if not headers:
            logger.error("获取认证请求头失败")
            return jsonify({
                'code': 401,
                'message': '获取认证信息失败',
                'timestamp': datetime.now().isoformat()
            }), 401
        
        logger.info(f"文件代理下载: {file_url}")
        
        # 调用内部服务器获取文件
        response = requests.get(
            file_url,
            headers=headers,
            timeout=60,
            stream=True
        )
        
        if response.status_code != 200:
            logger.error(f"内部服务返回错误: {response.status_code}")
            return jsonify({
                'code': response.status_code,
                'message': f'获取文件失败: {response.status_code}',
                'timestamp': datetime.now().isoformat()
            }), response.status_code
        
        # 获取文件名
        file_name = file_url.split('/')[-1]
        content_disposition = response.headers.get('Content-Disposition', '')
        if 'filename=' in content_disposition:
            import re
            match = re.search(r'filename=["\']?([^"\';]+)["\']?', content_disposition)
            if match:
                file_name = match.group(1)
        
        # 获取Content-Type
        content_type = response.headers.get('Content-Type', 'application/octet-stream')
        
        # 读取文件内容并返回（非流式，避免gunicorn worker问题）
        file_content = response.content
        
        from flask import Response
        from urllib.parse import quote
        
        # 对中文文件名进行URL编码，确保header格式正确
        encoded_filename = quote(file_name)
        
        return Response(
            file_content,
            status=200,
            headers={
                'Content-Type': content_type,
                'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_filename}",
                'Cache-Control': 'public, max-age=3600',
                'Content-Length': str(len(file_content))
            }
        )
        
    except requests.exceptions.Timeout:
        logger.error("文件下载超时")
        return jsonify({
            'code': 504,
            'message': '下载文件超时',
            'timestamp': datetime.now().isoformat()
        }), 504
    except requests.exceptions.RequestException as e:
        logger.error(f"文件下载请求异常: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'下载文件失败: {str(e)}',
            'timestamp': datetime.now().isoformat()
        }), 500
    except Exception as e:
        logger.error(f"文件代理错误: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器错误: {str(e)}',
            'timestamp': datetime.now().isoformat()
        }), 500


# ============================================
# 文档模板管理API
# ============================================

# 文档模板基础路径
DOC_TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), '文件生成')

def get_file_extension(filename):
    """获取文件扩展名"""
    return os.path.splitext(filename)[1].lower()

def remove_extension(filename):
    """移除文件扩展名"""
    return os.path.splitext(filename)[0]

def scan_directory_tree(base_path, relative_path=''):
    """
    递归扫描目录，返回树形结构
    最多三层：第一级是5个文件夹，第二级是文件或文件夹，第三级是文件
    """
    result = []
    current_path = os.path.join(base_path, relative_path)
    
    if not os.path.exists(current_path):
        return result
    
    def sort_key(item):
        """排序规则：不予受理文件夹排最前"""
        if item == '不予受理':
            return (0, '')  # 排最前
        elif os.path.isdir(os.path.join(current_path, item)):
            return (1, item)  # 其他文件夹其次，按名称排序
        else:
            return (2, item)  # 文件排最后，按名称排序
    
    try:
        items = sorted(os.listdir(current_path), key=sort_key)
        for item in items:
            # 跳过隐藏文件、临时文件和 output 文件夹
            if item.startswith('.') or item.startswith('~') or item.startswith('~$') or item == 'output':
                continue
                
            item_path = os.path.join(current_path, item)
            item_relative_path = os.path.join(relative_path, item) if relative_path else item
            
            if os.path.isdir(item_path):
                node = {
                    'name': item,
                    'displayName': item,
                    'type': 'folder',
                    'path': item_relative_path,
                    'children': scan_directory_tree(base_path, item_relative_path)
                }
                result.append(node)
            else:
                # 文件节点
                node = {
                    'name': item,
                    'displayName': remove_extension(item),
                    'type': 'file',
                    'path': item_relative_path,
                    'ext': get_file_extension(item)
                }
                result.append(node)
    except Exception as e:
        logger.error(f"扫描目录出错 {current_path}: {str(e)}")
    
    return result

@app.route('/api/doc_templates/tree', methods=['GET'])
def get_doc_templates_tree():
    """
    获取文档模板目录树结构
    返回文件生成文件夹的层级结构
    """
    try:
        if not os.path.exists(DOC_TEMPLATES_DIR):
            return jsonify({
                'code': 404,
                'message': '文档模板目录不存在',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 404
        
        tree = scan_directory_tree(DOC_TEMPLATES_DIR)
        
        return jsonify({
            'code': 200,
            'message': '获取成功',
            'data': tree,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logger.error(f"获取文档模板树失败: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器错误: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


@app.route('/api/doc_templates/download', methods=['GET'])
def download_doc_template():
    """
    下载文档模板文件
    参数: path - 文件相对路径
    """
    try:
        file_path = request.args.get('path', '')
        
        if not file_path:
            return jsonify({
                'code': 400,
                'message': '缺少参数: path',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 400
        
        # 安全检查：确保路径在文档模板目录内
        full_path = os.path.abspath(os.path.join(DOC_TEMPLATES_DIR, file_path))
        base_path = os.path.abspath(DOC_TEMPLATES_DIR)
        
        if not full_path.startswith(base_path):
            return jsonify({
                'code': 403,
                'message': '非法路径',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 403
        
        if not os.path.exists(full_path) or not os.path.isfile(full_path):
            return jsonify({
                'code': 404,
                'message': '文件不存在',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 404
        
        # 获取文件名和扩展名
        filename = os.path.basename(full_path)
        file_ext = os.path.splitext(filename)[1].lower()
        
        # 根据文件扩展名设置正确的 mimetype
        mimetype_map = {
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel'
        }
        mimetype = mimetype_map.get(file_ext)
        
        kwargs = {
            'directory': os.path.dirname(full_path),
            'path': os.path.basename(full_path),
            'as_attachment': True,
            'download_name': filename
        }
        if mimetype:
            kwargs['mimetype'] = mimetype
        
        return send_from_directory(**kwargs)
        
    except Exception as e:
        logger.error(f"下载文档模板失败: {str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器错误: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


# ============================================
# 文档生成API
# ============================================

@app.route('/api/doc_templates/generate', methods=['POST'])
def generate_document():
    """
    生成文档 - 根据模板和案件数据填充文档
    支持批量生成和合并
    参数:
      - template_path: 模板相对路径（单文件，向后兼容）
      - template_paths: 模板路径列表（多文件批量生成）
      - case_id: 案件ID（从立案详情API获取数据）
    返回:
      - 单文件: 直接返回文件
      - 多文件: 返回zip压缩包
    """
    try:
        # 获取请求参数
        params = request.get_json() or {}
        template_path = params.get('template_path', '')
        template_paths = params.get('template_paths', [])
        case_id = params.get('case_id', '')
        file_applicant_map = params.get('file_applicant_map', [])  # 文件与申请人的映射关系
        way = params.get('way', '')  # 结案方式：调解或裁决
        
        # 兼容单文件模式
        if template_path and not template_paths:
            template_paths = [template_path]
        
        if not template_paths:
            return jsonify({
                'code': 400,
                'message': '缺少参数: template_path 或 template_paths',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 400
        
        if not case_id:
            return jsonify({
                'code': 400,
                'message': '缺少参数: case_id',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 400
        
        # 安全检查所有模板路径
        logger.info(f"接收到的模板路径: {template_paths}, 数量: {len(template_paths)}")
        base_path = os.path.abspath(DOC_TEMPLATES_DIR)
        valid_paths = []
        for path in template_paths:
            full_path = os.path.abspath(os.path.join(DOC_TEMPLATES_DIR, path))
            if not full_path.startswith(base_path):
                return jsonify({
                    'code': 403,
                    'message': f'非法模板路径: {path}',
                    'data': None,
                    'timestamp': datetime.now().isoformat()
                }), 403
            if not os.path.exists(full_path):
                return jsonify({
                    'code': 404,
                    'message': f'模板文件不存在: {path}',
                    'data': None,
                    'timestamp': datetime.now().isoformat()
                }), 404
            # 检查文件格式
            file_ext = os.path.splitext(path)[1].lower()
            if file_ext not in ['.docx', '.xls', '.xlsx']:
                return jsonify({
                    'code': 400,
                    'message': f'不支持的文件格式: {file_ext}，请使用 .docx 或 .xlsx 格式',
                    'data': None,
                    'timestamp': datetime.now().isoformat()
                }), 400
            valid_paths.append(path)
        
        if not valid_paths:
            return jsonify({
                'code': 400,
                'message': '没有有效的模板文件',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 400
        
        # 调用内部API获取案件详情
        if not login_manager.check_and_renew_login():
            return jsonify({
                'code': 401,
                'message': '登录失败，无法查询案件详情',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 401
        
        headers = login_manager.get_auth_headers()
        if not headers:
            return jsonify({
                'code': 401,
                'message': '获取认证信息失败',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 401
        
        # 调用内部API获取案件详情
        detail_url = "http://10.96.10.78:8080/v1/api/admin/case/caseData"
        response = requests.post(
            detail_url,
            headers=headers,
            json={'id': case_id},
            timeout=30
        )
        
        if response.status_code != 200:
            return jsonify({
                'code': response.status_code,
                'message': f'获取案件详情失败: {response.status_code}',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 500
        
        case_data = response.json()
        
        # 处理API响应数据结构不一致的情况
        # 有时 data 是字典，有时可能是列表
        data_content = case_data.get('data')
        if isinstance(data_content, dict):
            case_no = data_content.get('case_no', case_id)
        elif isinstance(data_content, list) and len(data_content) > 0:
            # 如果是列表，取第一个元素的 case_no
            first_item = data_content[0]
            case_no = first_item.get('case_no', case_id) if isinstance(first_item, dict) else case_id
        else:
            case_no = case_id
        
        # 批量生成文档
        from batch_document_generator import BatchDocumentGenerator
        generator = BatchDocumentGenerator(DOC_TEMPLATES_DIR)
        result = generator.generate_batch(valid_paths, case_data, case_no, file_applicant_map=file_applicant_map, way=way)
        
        logger.info(f"批量文档生成成功: {result}")
        
        # 获取生成的文件列表和zip路径
        generated_files = result['zip']['files']
        logger.info(f"生成的文件数量: {len(generated_files)}, valid_paths数量: {len(valid_paths)}")
        zip_path = result['zip']['path']
        
        # 如果只有一个文件，直接返回该文件
        if len(generated_files) == 1:
            output_path = generated_files[0]
            output_filename = os.path.basename(output_path)
            file_ext = os.path.splitext(output_path)[1].lower()
            
            mimetype_map = {
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                '.xls': 'application/vnd.ms-excel'
            }
            mimetype = mimetype_map.get(file_ext)
            
            # 使用 send_file 并确保文件名正确编码
            from flask import send_file
            from urllib.parse import quote
            
            response = send_file(
                output_path,
                as_attachment=True,
                download_name=output_filename,
                mimetype=mimetype
            )
            
            # 手动设置 Content-Disposition 以确保文件名正确
            # 使用 RFC 5987 编码（filename*=UTF-8''）
            encoded_filename = quote(output_filename, safe='')
            response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
            
            return response
        
        # 多个文件，返回zip
        zip_filename = os.path.basename(zip_path)
        output_dir = os.path.dirname(zip_path)
        
        return send_from_directory(
            directory=output_dir,
            path=zip_filename,
            as_attachment=True,
            download_name=zip_filename,
            mimetype='application/zip'
        )
        
    except Exception as e:
        logger.error(f"生成文档失败: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            'code': 500,
            'message': f'生成文档失败: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


# ============================================
# 仲裁申请书生成API
# ============================================

@app.route('/api/application/generate', methods=['POST'])
def generate_application_document():
    """
    生成仲裁申请书Word文档
    前端传入拼接好的文本内容，后端替换模板占位符
    
    请求参数:
    {
        "applicant_info": "申请人：XXX，男...",  // 申请人信息段落
        "respondent_info": "被申请人：XXX...",   // 被申请人信息段落
        "requests": "1.裁决被申请人...\n2.裁决被申请人...",  // 仲裁请求
        "total_amount": "以上共计XXX元。",      // 总金额（可选）
        "facts_reasons": "申请人于...",         // 事实与理由
        "filename": "仲裁申请书-XXX.docx"       // 输出文件名（可选）
    }
    """
    try:
        # 获取请求参数
        params = request.get_json() or {}
        
        applicant_info = params.get('applicant_info', '')
        respondent_info = params.get('respondent_info', '')
        requests_text = params.get('requests', '')
        total_amount = params.get('total_amount', '')
        facts_reasons = params.get('facts_reasons', '')
        filename = params.get('filename', '仲裁申请书.docx')
        
        if not filename.endswith('.docx'):
            filename += '.docx'
        
        # 模板路径
        template_path = os.path.join('templates', '仲裁申请书模板.docx')
        if not os.path.exists(template_path):
            return jsonify({
                'code': 404,
                'message': '申请书模板文件不存在',
                'data': None,
                'timestamp': datetime.now().isoformat()
            }), 404
        
        # 输出目录和路径
        output_dir = os.path.join('文件生成', 'output')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, filename)
        
        # 加载模板并替换占位符
        from docx import Document
        doc = Document(template_path)
        
        # 准备替换数据
        data = {
            'applicant_info': applicant_info,
            'respondent_info': respondent_info,
            'requests': requests_text,
            'total_amount': total_amount,
            'facts_reasons': facts_reasons
        }
        
        # 替换段落中的占位符
        for para in doc.paragraphs:
            for key, value in data.items():
                placeholder = '{' + key + '}'
                if placeholder in para.text:
                    # 清除原有内容
                    para.clear()
                    
                    # 处理申请人和被申请人信息（需要加粗前缀）
                    if key in ['applicant_info', 'respondent_info']:
                        # 按行分割处理
                        lines = str(value).split('\n')
                        for i, line in enumerate(lines):
                            # 第一行已经有Word模板的首行缩进
                            # 从第二行开始需要手动添加缩进（两个全角空格）
                            if i > 0:
                                indent_run = para.add_run('　　')  # 两个全角空格
                                indent_run.font.name = '仿宋'
                                indent_run.font.size = Pt(15)
                                indent_run.font.bold = False
                                indent_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                            
                            # 找到第一个冒号的位置
                            colon_idx = line.find('：')
                            if colon_idx == -1:
                                colon_idx = line.find(':')
                            
                            if colon_idx > 0:
                                # 分割前缀和内容
                                prefix = line[:colon_idx + 1]  # 包含冒号
                                content = line[colon_idx + 1:]  # 冒号后的内容
                                
                                # 判断是否是需要加粗的前缀
                                # 申请人X： 和 被申请人X： 需要加粗
                                # 法定代表人： 不需要加粗
                                is_bold_prefix = False
                                if '申请人' in prefix and '法定代表人' not in prefix:
                                    is_bold_prefix = True
                                
                                # 前缀（根据条件决定是否加粗）
                                run = para.add_run(prefix)
                                run.font.name = '仿宋'
                                run.font.size = Pt(15)
                                run.font.bold = is_bold_prefix  # 根据条件加粗
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                                
                                # 内容不加粗
                                if content:
                                    run = para.add_run(content)
                                    run.font.name = '仿宋'
                                    run.font.size = Pt(15)
                                    run.font.bold = False
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                            else:
                                # 没有找到冒号，整行不加粗
                                run = para.add_run(line)
                                run.font.name = '仿宋'
                                run.font.size = Pt(15)
                                run.font.bold = False
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                            
                            # 添加换行（除了最后一行）
                            if i < len(lines) - 1:
                                para.add_run('\n')
                    elif key == 'requests':
                        # 请求事项需要处理换行缩进
                        lines = str(value).split('\n')
                        for i, line in enumerate(lines):
                            # 第一行已有Word模板首行缩进，从第二行开始手动添加缩进
                            if i > 0:
                                indent_run = para.add_run('　　')  # 两个全角空格
                                indent_run.font.name = '仿宋'
                                indent_run.font.size = Pt(15)
                                indent_run.font.bold = False
                                indent_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                            
                            run = para.add_run(line)
                            run.font.name = '仿宋'
                            run.font.size = Pt(15)
                            run.font.bold = False
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                            if i < len(lines) - 1:
                                para.add_run('\n')
                    elif key == 'facts_reasons':
                        # 事实与理由需要处理换行缩进
                        lines = str(value).split('\n')
                        for i, line in enumerate(lines):
                            # 第一行已有Word模板首行缩进，从第二行开始手动添加缩进
                            if i > 0:
                                indent_run = para.add_run('　　')  # 两个全角空格
                                indent_run.font.name = '仿宋'
                                indent_run.font.size = Pt(15)
                                indent_run.font.bold = False
                                indent_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                            
                            run = para.add_run(line)
                            run.font.name = '仿宋'
                            run.font.size = Pt(15)
                            run.font.bold = False
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                            if i < len(lines) - 1:
                                para.add_run('\n')
                    else:
                        # 其他内容正常处理（不加粗）
                        lines = str(value).split('\n')
                        for i, line in enumerate(lines):
                            run = para.add_run(line)
                            run.font.name = '仿宋'
                            run.font.size = Pt(15)
                            run.font.bold = False
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                            if i < len(lines) - 1:
                                para.add_run('\n')
        
        # 替换表格中的占位符（表格中不需要特殊加粗处理）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in data.items():
                            placeholder = '{' + key + '}'
                            if placeholder in para.text:
                                para.text = para.text.replace(placeholder, str(value))
                                for run in para.runs:
                                    run.font.name = '仿宋'
                                    run.font.size = Pt(15)
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        
        # 保存文档
        doc.save(output_path)
        logger.info(f"仲裁申请书生成成功: {output_path}")
        
        # 返回文件
        from flask import send_file
        from urllib.parse import quote
        
        response = send_file(
            output_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # 设置文件名编码
        encoded_filename = quote(filename, safe='')
        response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
        
        return response
        
    except Exception as e:
        logger.error(f"生成仲裁申请书失败: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            'code': 500,
            'message': f'生成仲裁申请书失败: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


# ============================================
# 裁决书制作API
# ============================================

@app.route('/api/award/elements/<case_id>', methods=['GET', 'POST'])
def award_elements(case_id):
    """
    获取或保存裁决书要素
    GET: 获取指定案件的裁决书要素
    POST: 保存裁决书要素
    """
    try:
        if request.method == 'GET':
            # 先获取案件详情，提取案号中的bianhao
            import re
            bianhao = case_id  # 默认使用case_id
            case_no = ''
            try:
                if login_manager.check_and_renew_login():
                    headers = login_manager.get_auth_headers()
                    detail_url = "http://10.96.10.78:8080/v1/api/admin/case/caseData"
                    response = requests.post(
                        detail_url,
                        headers=headers,
                        json={'id': case_id},
                        timeout=30
                    )
                    logger.info(f"[GET award/elements] case_id={case_id}, status={response.status_code}")
                    if response.status_code == 200:
                        case_data = response.json()
                        data_content = case_data.get('data')
                        if isinstance(data_content, dict):
                            case_no = data_content.get('case_no', '')
                        elif isinstance(data_content, list) and len(data_content) > 0:
                            first_item = data_content[0]
                            case_no = first_item.get('case_no', '') if isinstance(first_item, dict) else ''
                        
                        logger.info(f"[GET award/elements] case_no={case_no}")
                        
                        # 从案号提取bianhao
                        if case_no:
                            match = re.search(r'\[(\d{4})\](\d+)', case_no)
                            if match:
                                year, num = match.groups()
                                bianhao = f"{year}{num}"
                                logger.info(f"[GET award/elements] extracted bianhao={bianhao}")
                            else:
                                logger.warning(f"[GET award/elements] 无法从case_no提取bianhao: {case_no}")
            except Exception as e:
                logger.warning(f"[GET award/elements] 获取案件详情提取bianhao失败: {e}")
            
            logger.info(f"[GET award/elements] final bianhao={bianhao}, querying database...")
            
            # 查询数据库
            conn = db_manager.get_connection()
            cursor = conn.cursor(dictionary=True)
            
            try:
                cursor.execute(
                    "SELECT * FROM `裁决书要素保存` WHERE `案号` = %s",
                    (bianhao,)
                )
                result = cursor.fetchone()
                
                if result:
                    logger.info(f"[GET award/elements] 找到数据: 案号={bianhao}")
                    return jsonify({
                        'code': 200,
                        'message': 'success',
                        'data': result,
                        'timestamp': datetime.now().isoformat()
                    })
                else:
                    logger.info(f"[GET award/elements] 未找到数据: 案号={bianhao}")
                    return jsonify({
                        'code': 404,
                        'message': '未找到该案件的裁决书要素',
                        'data': None,
                        'timestamp': datetime.now().isoformat()
                    })
            finally:
                cursor.close()
                conn.close()
        
        else:  # POST
            import re
            data = request.get_json() or {}
            
            # 获取字段值
            case_no = data.get('case_no', case_id)
            arbitration_request = data.get('仲裁请求', '')
            applicant_claim = data.get('申请人称', '')
            respondent_claim = data.get('被申请人称', '')
            facts_found = data.get('经审理查明', '')
            committee_opinion = data.get('本委认为', '')
            final_decision = data.get('终局裁决', '')
            non_final_decision = data.get('非终局裁决', '')
            accept_time = data.get('受理时间', '')
            undisputed_facts = data.get('无争议事实', '')
            request_and_facts = data.get('仲裁请求和相关案件事实', '')
            
            # 转换受理时间：从中文格式(2025年3月2日)转换为紧凑格式(20250302)
            if accept_time:
                # 匹配 "2025年3月2日" 或 "2025年03月02日" 格式
                match = re.match(r'(\d{4})年(\d{1,2})月(\d{1,2})日', accept_time.strip())
                if match:
                    year, month, day = match.groups()
                    accept_time = f"{year}{int(month):02d}{int(day):02d}"
            
            # 从案号提取编号：明永劳人仲案字[2026]98号 -> 202698
            # 先尝试从case_no提取，如果没有则从case_id提取
            bianhao = case_id  # 默认使用case_id
            if case_no:
                match = re.search(r'\[(\d{4})\](\d+)', case_no)
                if match:
                    year, num = match.groups()
                    bianhao = f"{year}{num}"
                else:
                    # 尝试匹配 "明永劳人仲案字202698号" 这种格式
                    match = re.search(r'(\d{4})(\d+)', case_no)
                    if match:
                        year, num = match.groups()
                        bianhao = f"{year}{num}"
            
            # 连接数据库
            conn = db_manager.get_connection()
            cursor = conn.cursor()
            
            try:
                # 检查是否已存在（使用案号查询）
                cursor.execute(
                    "SELECT id FROM `裁决书要素保存` WHERE `案号` = %s",
                    (bianhao,)
                )
                existing = cursor.fetchone()
                
                if existing:
                    # 更新
                    cursor.execute("""
                        UPDATE `裁决书要素保存` SET
                            `仲裁请求` = %s,
                            `申请人称` = %s,
                            `被申请人称` = %s,
                            `经审理查明` = %s,
                            `本委认为` = %s,
                            `终局裁决` = %s,
                            `非终局裁决` = %s,
                            `受理时间` = %s,
                            `无争议事实` = %s,
                            `仲裁请求和相关案件事实` = %s
                        WHERE `案号` = %s
                    """, (
                        arbitration_request, applicant_claim, respondent_claim,
                        facts_found, committee_opinion, final_decision,
                        non_final_decision, accept_time, undisputed_facts,
                        request_and_facts, bianhao
                    ))
                else:
                    # 插入
                    cursor.execute("""
                        INSERT INTO `裁决书要素保存`
                        (`案号`, `仲裁请求`, `申请人称`, `被申请人称`, `经审理查明`,
                         `本委认为`, `终局裁决`, `非终局裁决`, `受理时间`,
                         `无争议事实`, `仲裁请求和相关案件事实`)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    """, (
                        bianhao, arbitration_request, applicant_claim, respondent_claim,
                        facts_found, committee_opinion, final_decision,
                        non_final_decision, accept_time, undisputed_facts,
                        request_and_facts
                    ))
                
                conn.commit()
                
                return jsonify({
                    'code': 200,
                    'message': '保存成功',
                    'data': {'case_id': case_id},
                    'timestamp': datetime.now().isoformat()
                })
                
            finally:
                cursor.close()
                conn.close()
                
    except Exception as e:
        logger.error(f"裁决书要素操作失败: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            'code': 500,
            'message': f'操作失败: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


@app.route('/api/award/generate', methods=['POST'])
def generate_award():
    """
    生成裁决书Word文档 - 使用长超时同步调用（参考申请人称模式）
    超时时间设为15分钟，避免异步状态同步问题
    """
    try:
        data = request.get_json() or {}
        case_id = data.get('case_id', '')
        case_no = data.get('case_no', '')
        
        # 获取庭审笔录
        if not login_manager.check_and_renew_login():
            return jsonify({'code': 401, 'message': '登录失败'}), 401
        
        headers = login_manager.get_auth_headers()
        response = requests.post(
            "http://10.96.10.78:8080/v1/api/admin/case/caseData",
            headers=headers,
            json={'id': case_id},
            timeout=30
        )
        
        if response.status_code != 200:
            return jsonify({'code': 500, 'message': '获取案件详情失败'}), 500
        
        case_data = response.json()
        data_content = case_data.get('data', {})
        if isinstance(data_content, list) and len(data_content) > 0:
            case_detail = data_content[0] if isinstance(data_content[0], dict) else {}
        elif isinstance(data_content, dict):
            case_detail = data_content
        else:
            case_detail = {}
        
        writing_json = case_detail.get('writing_json', []) if isinstance(case_detail, dict) else []
        
        # 提取笔录内容
        text_part1 = ''
        text_part2 = ''
        all_records_part3 = {}
        
        for item in writing_json:
            title = item.get('title', '')
            save_path = item.get('save_path', '')
            if '开庭笔录' in title or '庭审笔录' in title:
                json_str = item.get('json', '{}')
                try:
                    record_json = json.loads(json_str) if isinstance(json_str, str) else json_str
                    if '开庭笔录' in save_path or '庭审笔录' in save_path:
                        if not text_part1:
                            text_part1 = record_json.get('part1', '') or ''
                        if not text_part2:
                            text_part2 = record_json.get('part2', '') or ''
                    part3_content = record_json.get('part3', '') or ''
                    if part3_content:
                        date_match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', save_path)
                        if date_match:
                            year, month, day = date_match.groups()
                            formatted_date = f"{year}年{int(month)}月{int(day)}日"
                        else:
                            created_at = item.get('created_at', '')
                            date_match2 = re.search(r'(\d{4})-(\d{1,2})-(\d{1,2})', created_at)
                            if date_match2:
                                year, month, day = date_match2.groups()
                                formatted_date = f"{year}年{int(month)}月{int(day)}日"
                            else:
                                formatted_date = "未知日期"
                        
                        key = f"{title}_{formatted_date}撰写"
                        all_records_part3[key] = part3_content
                except:
                    pass
        
        text_part3 = json.dumps(all_records_part3, ensure_ascii=False) if all_records_part3 else ''
        
        if not text_part1 and not text_part2 and not text_part3:
            return jsonify({'code': 400, 'message': '未找到庭审笔录内容'}), 400
        
        # 提取案号编号
        bianhao = case_id
        if case_no:
            match = re.search(r'\[(\d{4})\](\d+)', case_no)
            if match:
                bianhao = f"{match.group(1)}{match.group(2)}"
        
        # 提取案件信息JSON
        handle_at_formatted = format_handle_at(case_detail.get('handle_at', ''))
        info_json = extract_court_record(text_part1, text_part2, text_part3, handle_at_formatted)
        info_json_str = json.dumps(info_json, ensure_ascii=False)
        
        logger.info(f"[GenerateAward] 开始调用Dify，预计最长等待15分钟...")
        
        # 调用Dify - 使用blocking模式，15分钟超时
        DIFY_API_KEY = "app-eEMlvxJweUDbvuOaJrUyaCeo"
        DIFY_BASE_URL = "http://127.0.0.1:8020/v1"
        
        payload = {
            "inputs": {
                "numb": bianhao,
                "textPart1": text_part1,
                "textPart2": text_part2,
                "textPart3": text_part3,
                "info": info_json_str
            },
            "response_mode": "blocking",
            "user": f"user-{case_id}"
        }
        
        # 设置15分钟超时（参考申请人称模式）
        resp = requests.post(
            f"{DIFY_BASE_URL}/workflows/run",
            headers={"Authorization": f"Bearer {DIFY_API_KEY}", "Content-Type": "application/json"},
            json=payload,
            timeout=900  # 15分钟
        )
        
        if resp.status_code != 200:
            error_text = resp.text[:500]
            logger.error(f"[GenerateAward] Dify错误: status={resp.status_code}, response={error_text}")
            return jsonify({'code': 500, 'message': f'Dify错误:{resp.status_code}, 详情:{error_text}'}), 500
        
        result = resp.json()
        logger.info(f"[GenerateAward] Dify返回: {result}")
        
        return jsonify({
            'code': 200,
            'message': '裁决书生成成功',
            'data': {'task_id': result.get('task_id', 'N/A')}
        })
            
    except requests.exceptions.Timeout:
        logger.error("[GenerateAward] Dify调用超时（超过15分钟）")
        return jsonify({'code': 504, 'message': '生成超时（超过15分钟），请稍后刷新页面查看结果'}), 504
    except Exception as e:
        logger.error(f"[GenerateAward] 异常: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'code': 500, 'message': str(e)}), 500


def get_case_material_files(case_id, auth_headers):
    """
    获取案件材料中的 word/docx 文件的 URL 列表
    从 arb/{id}/handle 接口获取 case_material（与 /api/handle/detail 一致）
    
    返回:
        list: 文件 URL 列表，格式为 [{"url": "xxx"}, ...]
              如果没有匹配文件则返回空列表
    """
    try:
        # 1. 从 arb/{id}/handle 接口获取 case_material
        # 注意：case_material 在 arb/{id}/handle 接口中，不在 case/caseData 中
        material_url = f"http://10.96.10.78:8080/v1/api/admin/arb/{case_id}/handle"
        material_resp = requests.get(material_url, headers=auth_headers, timeout=30)
        
        if material_resp.status_code != 200:
            logger.warning(f"获取案件材料失败: {material_resp.status_code}")
            return []
        
        material_data = material_resp.json()
        case_material = None
        
        # 尝试从不同层级获取 case_material（与 /api/handle/detail 逻辑一致）
        if isinstance(material_data, dict):
            if 'case_material' in material_data:
                case_material = material_data['case_material']
            elif 'data' in material_data and isinstance(material_data['data'], dict):
                if 'case_material' in material_data['data']:
                    case_material = material_data['data']['case_material']
                elif 'data' in material_data['data'] and isinstance(material_data['data']['data'], dict):
                    case_material = material_data['data']['data'].get('case_material')
        
        # 2. 筛选 file_path 包含 "word" 和 "docx" 的文件
        matched_files = []
        if isinstance(case_material, list):
            for item in case_material:
                file_path_str = item.get('file_path', '')
                if not file_path_str:
                    continue
                
                # 解析 file_path (JSON 字符串)
                try:
                    file_info_list = json.loads(file_path_str) if isinstance(file_path_str, str) else file_path_str
                    if isinstance(file_info_list, list):
                        for file_info in file_info_list:
                            url = file_info.get('url', '')
                            filename = file_info.get('filename', '')
                            # 检查是否包含 "word" 和 "docx"
                            if 'word' in url.lower() and 'docx' in url.lower():
                                matched_files.append({'url': url})
                            elif 'word' in filename.lower() and 'docx' in filename.lower():
                                matched_files.append({'url': url})
                except json.JSONDecodeError:
                    logger.warning(f"解析 file_path 失败: {file_path_str}")
                    continue
        
        logger.info(f"找到 {len(matched_files)} 个匹配的 word/docx 文件")
        return matched_files
        
    except Exception as e:
        logger.error(f"获取案件材料文件失败: {e}")
        return []


@app.route('/api/award/generate-draft', methods=['POST'])
def generate_award_draft():
    """
    一键生成裁决书初稿 - 使用长超时同步调用（参考申请人称模式）
    超时时间设为15分钟
    """
    try:
        data = request.get_json() or {}
        
        case_id = data.get('case_id', '')
        numb = data.get('numb', '')
        request_content = data.get('request', '')
        material = data.get('material', '')
        text_part3 = data.get('textPart3', '')
        slsj = data.get('slsj', '')
        
        if not case_id:
            return jsonify({'code': 400, 'message': '缺少案件ID'}), 400
        
        logger.info(f"[GenerateDraft] 收到请求: case_id={case_id}, numb={numb}")
        
        # 调用Dify
        DIFY_API_KEY = "app-OsDtggydgMq4R4NsqH111gBb"
        DIFY_BASE_URL = "http://127.0.0.1:8020/v1"
        DIFY_USER_ID = f"user-{case_id}-draft"
        
        # 获取认证头，用于获取案件材料文件
        auth_headers = login_manager.get_auth_headers()
        
        # 获取案件材料文件的 URL 列表
        matched_files = get_case_material_files(case_id, auth_headers)
        
        logger.info(f"[GenerateDraft] 找到 {len(matched_files)} 个匹配的 word/docx 文件")
        
        # 如果没有匹配的文件，使用默认 empty.docx 文件
        if not matched_files:
            default_file_url = "http://172.17.0.1:5000/files/empty.docx"
            matched_files = [{'url': default_file_url}]
            logger.info(f"[GenerateDraft] 使用默认 empty.docx")
        
        # 构建 files 参数（使用 remote_url 方式）
        files_param = []
        for file_info in matched_files:
            files_param.append({
                "transfer_method": "remote_url",
                "url": file_info['url'],
                "type": "document"
            })
        
        payload = {
            "inputs": {
                "numb": numb,
                "request": request_content,
                "material": material,
                "textPart3": text_part3,
                "slsj": slsj,
                "files": files_param
            },
            "response_mode": "blocking",
            "user": DIFY_USER_ID
        }
        
        logger.info(f"[GenerateDraft] 开始调用Dify，预计最长等待15分钟...")
        
        # 使用 blocking 模式，15分钟超时（参考申请人称模式）
        resp = requests.post(
            f"{DIFY_BASE_URL}/workflows/run",
            headers={"Authorization": f"Bearer {DIFY_API_KEY}", "Content-Type": "application/json"},
            json=payload,
            timeout=900  # 15分钟
        )
        
        if resp.status_code != 200:
            error_text = resp.text[:500]
            logger.error(f"[GenerateDraft] Dify错误: {resp.status_code}, {error_text}")
            return jsonify({'code': 500, 'message': f'Dify错误:{resp.status_code}'}), 500
        
        result = resp.json()
        logger.info(f"[GenerateDraft] Dify返回: {result}")
        
        # 从返回结果中提取
        result_data = result.get('data', {})
        outputs = result_data.get('outputs', {})
        
        # Dify生成成功后，从数据库查询裁决书要素并返回
        elements_data = {}
        try:
            # 从numb构建案号（如202688）
            conn = db_manager.get_connection()
            cursor = conn.cursor(dictionary=True)
            try:
                cursor.execute(
                    "SELECT * FROM `裁决书要素保存` WHERE `案号` = %s",
                    (numb,)
                )
                db_result = cursor.fetchone()
                if db_result:
                    elements_data = {
                        '仲裁请求': db_result.get('仲裁请求', ''),
                        '申请人称': db_result.get('申请人称', ''),
                        '被申请人称': db_result.get('被申请人称', ''),
                        '经审理查明': db_result.get('经审理查明', ''),
                        '本委认为': db_result.get('本委认为', ''),
                        '终局裁决': db_result.get('终局裁决', ''),
                        '非终局裁决': db_result.get('非终局裁决', '')
                    }
                    logger.info(f"[GenerateDraft] 从数据库获取裁决书要素成功: {numb}")
                else:
                    logger.warning(f"[GenerateDraft] 数据库中未找到裁决书要素: {numb}")
            finally:
                cursor.close()
                conn.close()
        except Exception as e:
            logger.warning(f"[GenerateDraft] 查询数据库失败: {e}")
        
        return jsonify({
            'code': 200,
            'message': '✅ 初稿生成成功！',
            'data': {
                '无争议事实': outputs.get('无争议事实', ''),
                '仲裁请求和相关案件事实': outputs.get('仲裁请求和相关案件事实', ''),
                'elements': elements_data
            }
        })
            
    except requests.exceptions.Timeout:
        logger.error("[GenerateDraft] Dify调用超时（超过15分钟）")
        return jsonify({'code': 504, 'message': '生成超时（超过15分钟），请稍后刷新页面查看结果'}), 504
    except Exception as e:
        logger.error(f"[GenerateDraft] 异常: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'code': 500, 'message': str(e)}), 500


@app.route('/api/workflow/generate-claim', methods=['POST'])
def generate_claim_workflow():
    """
    调用Dify Workflow自动生成申请人称或被申请人称
    调用URL和生成Word一致，但key不同：
    - 申请人称: app-S5FnZhPbetGmnw6A8mefsCVj
    - 被申请人称: app-iPP2ZBN20yYiYC6hT8pu04Rj
    
    请求参数:
    - case_id: 案件ID
    - case_no: 案号
    - claim_type: 'applicantClaim' 或 'respondentClaim'
    - count: 限制字数
    - request: 仲裁请求内容
    - textPart3: 庭审笔录part3内容
    - is_applicant: 是否为申请人称
    """
    try:
        data = request.get_json() or {}
        
        case_id = data.get('case_id', '')
        case_no = data.get('case_no', '')
        claim_type = data.get('claim_type', '')
        count = data.get('count', 600)
        request_content = data.get('request', '')
        text_part3 = data.get('textPart3', '')
        is_applicant = data.get('is_applicant', True)
        
        logger.info(f"[GenerateClaim] 接收到请求: claim_type={claim_type}, is_applicant={is_applicant}")
        
        # Dify 配置
        # 根据类型选择不同的key
        if claim_type == 'respondentClaim':
            DIFY_API_KEY = "app-iPP2ZBN20yYiYC6hT8pu04Rj"  # 被申请人称
        else:
            DIFY_API_KEY = "app-S5FnZhPbetGmnw6A8mefsCVj"  # 申请人称
        DIFY_BASE_URL = "http://127.0.0.1:8020/v1"
        DIFY_USER_ID = f"user-{case_id}-claim" if case_id else "user-claim"
        
        logger.info(f"[GenerateClaim] 调用Workflow, case_id={case_id}, claim_type={claim_type}, count={count}")
        logger.info(f"[GenerateClaim] 使用API Key: {DIFY_API_KEY[:20]}...")
        logger.info(f"[GenerateClaim] 参数: count={count}, request长度={len(request_content)}, textPart3长度={len(text_part3)}")
        
        # 调用 Dify Workflow（阻塞模式，等待结果）
        workflow_url = f"{DIFY_BASE_URL}/workflows/run"
        workflow_headers = {
            "Authorization": f"Bearer {DIFY_API_KEY}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "inputs": {
                "count": str(count),
                "request": request_content,
                "textPart3": text_part3
            },
            "response_mode": "blocking",
            "user": DIFY_USER_ID
        }
        
        # 设置较长的超时时间（约1分钟）
        workflow_resp = requests.post(workflow_url, headers=workflow_headers, json=payload, timeout=120)
        
        if workflow_resp.status_code != 200:
            workflow_result = workflow_resp.json() if workflow_resp.text else {}
            logger.error(f"[GenerateClaim] Workflow调用失败: {workflow_result}")
            return jsonify({
                'success': False,
                'content': None,
                'message': f'Dify Workflow调用失败: {workflow_result.get("message", "未知错误")}'
            }), 500
        
        workflow_result = workflow_resp.json()
        logger.info(f"[GenerateClaim] Workflow返回: {workflow_result}")
        
        # 从返回结果中提取生成的内容
        # Dify Workflow 返回结构可能需要根据实际情况调整
        result_data = workflow_result.get('data', {})
        
        # 尝试从outputs中获取结果
        outputs = result_data.get('outputs', {})
        generated_content = outputs.get('result', '') or outputs.get('content', '') or outputs.get('text', '')
        
        # 如果outputs中没有，尝试从answer或其他字段获取
        if not generated_content:
            generated_content = result_data.get('answer', '') or result_data.get('content', '')
        
        if not generated_content:
            return jsonify({
                'success': False,
                'content': None,
                'message': '工作流返回结果为空'
            }), 500
        
        return jsonify({
            'success': True,
            'content': generated_content,
            'message': '生成成功'
        })
        
    except requests.exceptions.Timeout:
        logger.error("[GenerateClaim] Workflow调用超时")
        return jsonify({
            'success': False,
            'content': None,
            'message': '工作流调用超时（超过2分钟），请稍后重试'
        }), 504
    except Exception as e:
        logger.error(f"生成称述失败: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'content': None,
            'message': f'生成失败: {str(e)}'
        }), 500


@app.route('/api/holiday/list', methods=['GET'])
def get_holiday_list():
    """
    获取休假日列表
    参数: yearMonth (格式: YYYYMM, 如 202602)
    """
    try:
        year_month = request.args.get('yearMonth', '')
        if not year_month or len(year_month) != 6 or not year_month.isdigit():
            return jsonify({
                'success': False,
                'message': '参数错误: yearMonth 格式应为 YYYYMM (如 202602)'
            }), 400
        
        global _holiday_cache, _holiday_cache_month
        
        # 检查缓存
        if _holiday_cache_month == year_month and year_month in _holiday_cache:
            return jsonify({
                'success': True,
                'message': '查询成功(缓存)',
                'data': _holiday_cache[year_month]
            })
        
        # 检查登录状态
        if not login_manager.check_and_renew_login():
            return jsonify({
                'success': False,
                'message': '登录失败，无法查询休假日'
            }), 401
        
        headers = login_manager.get_auth_headers()
        if not headers:
            return jsonify({
                'success': False,
                'message': '获取认证信息失败'
            }), 401
        
        url = f"http://10.96.10.78:8080/v1/api/admin/holidayList?date={year_month}"
        
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        result = response.json()
        
        if result.get('code') == 200:
            # 解析休假日数据
            holidays = result.get('data', {}).get('holiaya', [])
            # 提取日期中的日份 (20260223 -> 23)
            holiday_days = []
            for h in holidays:
                date_str = h.get('date', '')
                if len(date_str) == 8:
                    try:
                        day = int(date_str[6:8])
                        holiday_days.append(day)
                    except:
                        pass
            
            # 更新缓存
            _holiday_cache = { year_month: holiday_days }
            _holiday_cache_month = year_month
            
            return jsonify({
                'success': True,
                'message': '查询成功',
                'data': holiday_days
            })
        else:
            return jsonify({
                'success': False,
                'message': result.get('message', '查询休假日失败')
            }), 500
            
    except requests.exceptions.RequestException as e:
        logger.error(f"休假日查询请求失败: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'查询请求失败: {str(e)}'
        }), 500
    except Exception as e:
        logger.error(f"休假日查询接口错误: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'服务器错误: {str(e)}'
        }), 500


@app.route('/api/internship/sign', methods=['POST'])
def generate_internship_sign():
    """
    生成见习签到表（Excel格式）
    按签到表.xlsx模板生成
    """
    try:
        global _holiday_cache, _holiday_cache_month
        data = request.get_json()
        
        if not data:
            return jsonify({
                'success': False,
                'message': '请求数据为空'
            }), 400
        
        year = data.get('year')
        month = data.get('month')
        persons = data.get('persons', [])
        
        if not year or not month:
            return jsonify({
                'success': False,
                'message': '请选择年份和月份'
            }), 400
        
        if not persons or len(persons) == 0:
            return jsonify({
                'success': False,
                'message': '请至少添加一名见习人员'
            }), 400
        
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        from openpyxl.utils import get_column_letter
        from datetime import datetime, timedelta
        import calendar
        
        wb = Workbook()
        ws = wb.active
        ws.title = f"{year}年{month}月"
        
        # 样式定义
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        center_align = Alignment(horizontal='center', vertical='center')
        
        # 字体
        font_title = Font(name='宋体', size=14, bold=True)
        font_normal = Font(name='宋体', size=11)
        font_header = Font(name='宋体', size=11, bold=True)
        
        # 列宽
        ws.column_dimensions['A'].width = 11.625
        ws.column_dimensions['B'].width = 11.625
        ws.column_dimensions['C'].width = 11.625
        for col in ['D', 'E', 'F', 'G', 'H', 'I', 'J', 'K']:
            ws.column_dimensions[col].width = 9
        
        # ===== 标题行 =====
        ws.merge_cells('A1:K1')
        cell = ws['A1']
        cell.value = f'{year}年{month}月永安市社会保险中心见习人员签到表'
        cell.font = font_title
        cell.alignment = center_align
        for col in range(1, 12):
            ws.cell(row=1, column=col).border = thin_border
        ws.row_dimensions[1].height = 22.5
        
        # ===== 表头 =====
        # 第2行
        headers_row2 = ['日期', '星期', '姓名', '上午', '', '', '', '下午', '', '', '']
        for col, val in enumerate(headers_row2, 1):
            cell = ws.cell(row=2, column=col, value=val)
            cell.font = font_header
            cell.alignment = center_align
            cell.border = thin_border
        ws.row_dimensions[2].height = 23.1
        
        # 第3行
        headers_row3 = ['', '', '', '上班', '', '下班', '', '上班', '', '下班', '']
        for col, val in enumerate(headers_row3, 1):
            cell = ws.cell(row=3, column=col, value=val)
            cell.font = font_header
            cell.alignment = center_align
            cell.border = thin_border
        ws.row_dimensions[3].height = 23.1
        
        # 第4行
        headers_row4 = ['', '', '', '时间', '签名', '时间', '签名', '时间', '签名', '时间', '签名']
        for col, val in enumerate(headers_row4, 1):
            cell = ws.cell(row=4, column=col, value=val)
            cell.font = font_header
            cell.alignment = center_align
            cell.border = thin_border
        ws.row_dimensions[4].height = 23.1
        
        # 合并单元格
        ws.merge_cells('A2:A4')  # 日期
        ws.merge_cells('B2:B4')  # 星期
        ws.merge_cells('C2:C4')  # 姓名
        ws.merge_cells('D2:G2')  # 上午
        ws.merge_cells('H2:K2')  # 下午
        ws.merge_cells('D3:E3')  # 上班
        ws.merge_cells('F3:G3')  # 下班
        ws.merge_cells('H3:I3')  # 上班
        ws.merge_cells('J3:K3')  # 下班
        
        # ===== 获取休假日 =====
        holiday_days = set()
        try:
            year_month = f"{year}{str(month).zfill(2)}"
            if _holiday_cache_month == year_month and year_month in _holiday_cache:
                holiday_days = set(_holiday_cache[year_month])
            else:
                if login_manager.check_and_renew_login():
                    headers = login_manager.get_auth_headers()
                    if headers:
                        resp = requests.get(
                            f"http://10.96.10.78:8080/v1/api/admin/holidayList?date={year_month}",
                            headers=headers, timeout=30
                        )
                        if resp.status_code == 200:
                            result = resp.json()
                            if result.get('code') == 200:
                                for h in result.get('data', {}).get('holiaya', []):
                                    ds = h.get('date', '')
                                    if len(ds) == 8:
                                        try:
                                            holiday_days.add(int(ds[6:8]))
                                        except:
                                            pass
                                _holiday_cache[year_month] = list(holiday_days)
                                _holiday_cache_month = year_month
        except Exception as e:
            logger.warning(f"获取休假日失败: {e}")
        
        # ===== 填充数据 =====
        current_row = 5
        last_day = calendar.monthrange(year, month)[1]
        
        for day in range(1, last_day + 1):
            date_obj = datetime(year, month, day)
            weekday = date_obj.weekday()  # 0=周一, 5=周六, 6=周日
            
            # 跳过周末和节假日
            if weekday >= 5 or day in holiday_days:
                continue
            
            weekday_names = ['星期一', '星期二', '星期三', '星期四', '星期五', '星期六', '星期日']
            weekday_str = weekday_names[weekday]
            
            # 找出在该日期范围内的见习人员
            active_persons = []
            for p in persons:
                try:
                    start = datetime.strptime(p.get('startDate', ''), '%Y-%m-%d')
                    end = datetime.strptime(p.get('endDate', ''), '%Y-%m-%d')
                    if start <= date_obj <= end:
                        active_persons.append(p.get('name', ''))
                except:
                    pass
            
            if len(active_persons) == 0:
                continue
            
            # 第一个人员行：显示日期
            for pi, name in enumerate(active_persons):
                row = current_row + pi
                ws.row_dimensions[row].height = 23.1
                
                # 日期列（只第一行显示）
                if pi == 0:
                    cell = ws.cell(row=row, column=1)
                    # 使用1900年的日期，格式m/d显示月/日
                    cell.value = datetime(1900 if day > 1 else year, month, day)
                    cell.number_format = 'm/d;@'
                    cell.font = font_normal
                    cell.alignment = center_align
                    cell.border = thin_border
                    
                    # 星期列
                    cell = ws.cell(row=row, column=2, value=weekday_str)
                    cell.font = font_normal
                    cell.alignment = center_align
                    cell.border = thin_border
                else:
                    # 后续人员行：日期和星期为空但保留边框
                    ws.cell(row=row, column=1).border = thin_border
                    ws.cell(row=row, column=2).border = thin_border
                
                # 姓名
                cell = ws.cell(row=row, column=3, value=name)
                cell.font = font_normal
                cell.alignment = center_align
                cell.border = thin_border
                
                # 上午下午各列（空白）
                for col in range(4, 12):
                    ws.cell(row=row, column=col).border = thin_border
            
            current_row += len(active_persons)
        
        # 保存文件
        output_dir = '文件生成/output'
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"见习签到表_{year}年{month}月.xlsx"
        filepath = os.path.join(output_dir, filename)
        
        wb.save(filepath)
        
        logger.info(f"生成见习签到表: {filepath}")
        
        return jsonify({
            'success': True,
            'message': '生成成功',
            'data': {
                'filename': filename,
                'download_url': f'/api/doc_templates/download?path={output_dir}/{filename}'
            }
        })
        
    except Exception as e:
        logger.error(f"生成见习签到表失败: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'生成失败: {str(e)}'
        }), 500


@app.route('/api/workflow/optimize-text', methods=['POST'])
def optimize_text_workflow():
    """
    调用Dify Workflow一键优化文本（申请人称/被申请人称/经审理查明）
    Dify Key: app-YjTFrQ3LKrFyK5Q7c0CajCJN
    
    请求参数:
    - type: 1=申请人称, 2=被申请人称, 3=经审理查明
    - numb: 8位数案号
    - text: 需要优化的原始文本内容
    - user_request: 用户的优化建议（可选）
    """
    try:
        data = request.get_json() or {}
        
        type_code = data.get('type', 1)
        numb = data.get('numb', '')
        text = data.get('text', '')
        user_request = data.get('user_request', '')
        
        # 参数校验
        if not text or not text.strip():
            logger.warning(f"[OptimizeText] 文本内容为空")
            return jsonify({
                'success': False,
                'content': None,
                'message': '文本内容不能为空'
            }), 400
        
        if not numb:
            logger.warning(f"[OptimizeText] 案号为空")
            return jsonify({
                'success': False,
                'content': None,
                'message': '案号不能为空'
            }), 400
        
        # 类型映射
        type_names = {1: '申请人称', 2: '被申请人称', 3: '经审理查明'}
        type_name = type_names.get(type_code, '未知类型')
        
        logger.info(f"[OptimizeText] 接收到请求: type={type_code}({type_name}), numb={numb}, text长度={len(text)}, user_request={user_request[:50] if user_request else '无'}...")
        logger.info(f"[OptimizeText] text前100字: {text[:100]}...")
        
        # Dify 配置
        DIFY_API_KEY = "app-YjTFrQ3LKrFyK5Q7c0CajCJN"
        DIFY_BASE_URL = "http://127.0.0.1:8020/v1"
        DIFY_USER_ID = f"user-{numb}-optimize-{type_code}"
        
        # 调用 Dify Workflow（阻塞模式，等待结果）
        workflow_url = f"{DIFY_BASE_URL}/workflows/run"
        workflow_headers = {
            "Authorization": f"Bearer {DIFY_API_KEY}",
            "Content-Type": "application/json"
        }
        
        # 如果用户没有提供优化建议，使用"无"
        if not user_request or not user_request.strip():
            user_request = "无"
        
        payload = {
            "inputs": {
                "numb": str(numb),
                "type": str(type_code),
                "text": text,
                "userRequest": user_request
            },
            "response_mode": "blocking",
            "user": DIFY_USER_ID
        }
        
        logger.info(f"[OptimizeText] 调用Workflow: type={type_code}, numb={numb}")
        logger.info(f"[OptimizeText] 请求payload: {payload}")
        
        # 设置较长的超时时间（约2分钟）
        workflow_resp = requests.post(workflow_url, headers=workflow_headers, json=payload, timeout=120)
        
        logger.info(f"[OptimizeText] Workflow响应状态码: {workflow_resp.status_code}")
        
        if workflow_resp.status_code != 200:
            try:
                workflow_result = workflow_resp.json() if workflow_resp.text else {}
            except:
                workflow_result = {'raw_response': workflow_resp.text[:500]}
            logger.error(f"[OptimizeText] Workflow调用失败: {workflow_result}")
            return jsonify({
                'success': False,
                'content': None,
                'message': f'优化服务调用失败: {workflow_result.get("message", "未知错误")}'
            }), 500
        
        try:
            workflow_result = workflow_resp.json()
            logger.info(f"[OptimizeText] Workflow返回: {workflow_result}")
        except Exception as e:
            logger.error(f"[OptimizeText] 解析Workflow响应失败: {str(e)}, 原始响应: {workflow_resp.text[:500]}")
            return jsonify({
                'success': False,
                'content': None,
                'message': f'解析优化服务响应失败: {str(e)}'
            }), 500
        
        # 从返回结果中提取优化后的内容
        result_data = workflow_result.get('data', {})
        outputs = result_data.get('outputs', {})
        
        # 尝试从outputs中获取结果
        optimized_content = outputs.get('result', '') or outputs.get('content', '') or outputs.get('text', '')
        
        # 如果outputs中没有，尝试从answer或其他字段获取
        if not optimized_content:
            optimized_content = result_data.get('answer', '') or result_data.get('content', '')
        
        if not optimized_content:
            return jsonify({
                'success': False,
                'content': None,
                'message': '优化服务返回结果为空'
            }), 500
        
        logger.info(f"[OptimizeText] 优化成功: type={type_code}, 输出长度={len(optimized_content)}")
        
        return jsonify({
            'success': True,
            'content': optimized_content,
            'message': '优化成功'
        })
        
    except requests.exceptions.Timeout:
        logger.error("[OptimizeText] Workflow调用超时")
        return jsonify({
            'success': False,
            'content': None,
            'message': '优化服务调用超时（超过2分钟），请稍后重试'
        }), 504
    except Exception as e:
        logger.error(f"优化文本失败: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'content': None,
            'message': f'优化失败: {str(e)}'
        }), 500


def build_party_role_map(case_detail):
    """
    根据案件详情构建当事人role映射
    返回: {submitter_name: role}
    """
    role_map = {}
    
    # 处理申请人
    applicant_arr = case_detail.get('applicant_arr', []) if isinstance(case_detail, dict) else []
    if isinstance(applicant_arr, list):
        if len(applicant_arr) == 1:
            role_map[applicant_arr[0].get('name', '')] = "申请人"
        else:
            for idx, app in enumerate(applicant_arr, 1):
                role_map[app.get('name', '')] = f"申请人{idx}"
    
    # 处理被申请人
    respondent_arr = case_detail.get('respondent_arr', []) if isinstance(case_detail, dict) else []
    if isinstance(respondent_arr, list):
        if len(respondent_arr) == 1:
            role_map[respondent_arr[0].get('name', '')] = "被申请人"
        else:
            for idx, resp in enumerate(respondent_arr, 1):
                role_map[resp.get('name', '')] = f"被申请人{idx}"
    
    # 处理第三人（支持 thirdparty_arr 和 thirdpartys_arr 两种字段名）
    thirdparty_arr = case_detail.get('thirdparty_arr', []) if isinstance(case_detail, dict) else []
    thirdpartys_arr = case_detail.get('thirdpartys_arr', []) if isinstance(case_detail, dict) else []
    
    # 优先使用 thirdpartys_arr，如果不存在则使用 thirdparty_arr
    third_parties = thirdpartys_arr if thirdpartys_arr else thirdparty_arr
    
    if isinstance(third_parties, list):
        if len(third_parties) == 1:
            role_map[third_parties[0].get('name', '')] = "第三人"
        else:
            for idx, third in enumerate(third_parties, 1):
                role_map[third.get('name', '')] = f"第三人{idx}"
    
    return role_map


def build_evidence_files_and_mapping(case_evidence_material, case_detail=None):
    """
    构建证据文件列表和文件名映射
    
    返回:
        tuple: (evidence_files, filename_mapping)
        - evidence_files: [{"transfer_method": "remote_url", "url": "原始URL", "type": "document"}, ...]
        - filename_mapping: [{"raw_name": "原始文件名.pdf", "display_name": "重命名后的文件名"}, ...]
    """
    evidence_files = []
    filename_mapping = []
    
    # 构建当事人role映射
    role_map = build_party_role_map(case_detail) if case_detail else {}
    logger.info(f"[BuildEvidence] 当事人role映射: {role_map}")
    
    if not isinstance(case_evidence_material, list):
        logger.warning(f"[BuildEvidence] case_evidence_material 不是列表: {type(case_evidence_material)}")
        return evidence_files, filename_mapping
    
    logger.info(f"[BuildEvidence] 处理 {len(case_evidence_material)} 条证据材料记录")
    
    # 用于记录每个role的证据序号计数器（基于原始顺序，非PDF也计数）
    role_counter = {}
    
    for idx, item in enumerate(case_evidence_material):
        file_path_str = item.get('file_path', '')
        if not file_path_str:
            continue
        
        # 获取submitter并计算序号（在检查PDF之前，确保所有证据都计数）
        submitter = item.get('submitter', '未知提交人')
        role = role_map.get(submitter, submitter)
        role_counter[role] = role_counter.get(role, 0) + 1
        no = role_counter[role]
        
        # 检查是否为PDF文件
        if '.pdf' not in file_path_str.lower():
            continue
        
        try:
            # 解析 file_path
            file_info_list = json.loads(file_path_str) if isinstance(file_path_str, str) else file_path_str
            if not isinstance(file_info_list, list) or len(file_info_list) == 0:
                continue
            
            file_info = file_info_list[0]
            url = file_info.get('url', '')
            
            if not url:
                continue
            
            # 构建重命名后的文件名
            name = item.get('name', '未知名称')
            material_type = item.get('material_type', '未知类型')
            object_name = item.get('object', '未知对象')
            evidence_type = item.get('type', '')  # 证据类型，如"书证"
            
            def clean_filename(s):
                return re.sub(r'[\\/:*?"<>|]', '_', str(s))
            
            # 从name中提取序号（格式: {名称}_{序号}）
            # 如果name中有下划线，取最后一部分作为序号
            if '_' in name:
                parts = name.rsplit('_', 1)
                name_without_number = parts[0]
                number = parts[1]
            else:
                name_without_number = name
                number = "未知"
            
            # 文件名格式: {role}_{名称}_{序号}_{material_type}_{object}_{type}（无后缀）
            if evidence_type:
                display_name = f"{clean_filename(role)}_{clean_filename(name_without_number)}_{clean_filename(number)}_{clean_filename(material_type)}_{clean_filename(object_name)}_{clean_filename(evidence_type)}"
            else:
                display_name = f"{clean_filename(role)}_{clean_filename(name_without_number)}_{clean_filename(number)}_{clean_filename(material_type)}_{clean_filename(object_name)}"
            
            logger.info(f"[BuildEvidence] 添加文件: {display_name} -> {url[:80]}...")
            
            # 使用原始URL（10.96.10.78:8080）
            evidence_files.append({
                "transfer_method": "remote_url",
                "url": url,
                "type": "document"
            })
            
            # 从URL中提取原始文件名
            raw_name = url.split('/')[-1] if '/' in url else 'unknown.pdf'
            
            # 记录映射关系（原始文件名、重命名后的文件名和完整URL）
            filename_mapping.append({
                "raw_name": raw_name,
                "display_name": display_name,
                "url": url
            })
            
        except Exception as e:
            logger.warning(f"[BuildEvidence] 处理证据文件失败: {e}")
            continue
    
    logger.info(f"[BuildEvidence] 共处理 {len(evidence_files)} 个PDF证据文件")
    return evidence_files, filename_mapping


@app.route('/api/workflow/analyze-evidence', methods=['POST'])
def analyze_evidence():
    """
    证据分析工作流
    调用Dify Workflow分析案件证据材料
    
    Dify Key: app-z6fLFQnv0c9VFnz9LtX0gWt5
    参数:
    - numb: 案号编号（如2025432）
    - textPart3: 庭审笔录内容
    - evidence: PDF证据文件列表（下载后重命名，通过本地URL访问）
    """
    try:
        data = request.get_json() or {}
        case_id = data.get('case_id', '')
        case_no = data.get('case_no', '')
        
        if not case_id:
            return jsonify({'code': 400, 'message': '缺少案件ID'}), 400
        
        logger.info(f"[AnalyzeEvidence] 收到请求: case_id={case_id}, case_no={case_no}")
        
        # 获取认证头
        if not login_manager.check_and_renew_login():
            return jsonify({'code': 401, 'message': '登录失败'}), 401
        
        headers = login_manager.get_auth_headers()
        
        # ========== 1. 获取庭审笔录 textPart3 ==========
        response = requests.post(
            "http://10.96.10.78:8080/v1/api/admin/case/caseData",
            headers=headers,
            json={'id': case_id},
            timeout=30
        )
        
        if response.status_code != 200:
            return jsonify({'code': 500, 'message': '获取案件详情失败'}), 500
        
        case_data = response.json()
        data_content = case_data.get('data', {})
        if isinstance(data_content, list) and len(data_content) > 0:
            case_detail = data_content[0] if isinstance(data_content[0], dict) else {}
        elif isinstance(data_content, dict):
            case_detail = data_content
        else:
            case_detail = {}
        
        writing_json = case_detail.get('writing_json', []) if isinstance(case_detail, dict) else []
        
        # 提取textPart3（和生成word一样的逻辑）
        all_records_part3 = {}
        for item in writing_json:
            title = item.get('title', '')
            save_path = item.get('save_path', '')
            if '开庭笔录' in title or '庭审笔录' in title:
                json_str = item.get('json', '{}')
                try:
                    record_json = json.loads(json_str) if isinstance(json_str, str) else json_str
                    part3_content = record_json.get('part3', '') or ''
                    if part3_content:
                        # 提取日期
                        date_match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', save_path)
                        if date_match:
                            year, month, day = date_match.groups()
                            formatted_date = f"{year}年{int(month)}月{int(day)}日"
                        else:
                            created_at = item.get('created_at', '')
                            date_match2 = re.search(r'(\d{4})-(\d{1,2})-(\d{1,2})', created_at)
                            if date_match2:
                                year, month, day = date_match2.groups()
                                formatted_date = f"{year}年{int(month)}月{int(day)}日"
                            else:
                                formatted_date = "未知日期"
                        
                        key = f"{title}_{formatted_date}撰写"
                        all_records_part3[key] = part3_content
                except:
                    pass
        
        text_part3 = json.dumps(all_records_part3, ensure_ascii=False) if all_records_part3 else ''
        
        if not text_part3:
            return jsonify({'code': 400, 'message': '未找到庭审笔录内容'}), 400
        
        # ========== 2. 提取案号编号 numb ==========
        bianhao = case_id
        if case_no:
            match = re.search(r'\[(\d{4})\](\d+)', case_no)
            if match:
                bianhao = f"{match.group(1)}{match.group(2)}"
        
        logger.info(f"[AnalyzeEvidence] numb={bianhao}")
        
        # ========== 3. 获取并下载证据材料 PDF 文件 ==========
        # 从 arb/{id}/handle 接口获取 case_evidence_material
        material_url = f"http://10.96.10.78:8080/v1/api/admin/arb/{case_id}/handle"
        material_resp = requests.get(material_url, headers=headers, timeout=30)
        
        case_evidence_material = None
        
        if material_resp.status_code == 200:
            material_data = material_resp.json()
            
            # 尝试从不同层级获取 case_evidence_material
            if isinstance(material_data, dict):
                if 'case_evidence_material' in material_data:
                    case_evidence_material = material_data['case_evidence_material']
                elif 'data' in material_data and isinstance(material_data['data'], dict):
                    if 'case_evidence_material' in material_data['data']:
                        case_evidence_material = material_data['data']['case_evidence_material']
                    elif 'data' in material_data['data'] and isinstance(material_data['data']['data'], dict):
                        case_evidence_material = material_data['data']['data'].get('case_evidence_material')
        
        # 构建证据文件列表和文件名映射（使用原始URL，不再下载）
        evidence_files, filename_mapping = build_evidence_files_and_mapping(
            case_evidence_material,
            case_detail=case_detail
        )
        
        if not evidence_files:
            return jsonify({'code': 400, 'message': '未找到PDF格式的证据材料'}), 400
        
        # ========== 4. 调用Dify Workflow ==========
        DIFY_API_KEY = "app-z6fLFQnv0c9VFnz9LtX0gWt5"
        DIFY_BASE_URL = "http://127.0.0.1:8020/v1"
        DIFY_USER_ID = f"user-{case_id}-evidence"
        
        # filename_mapping 转为 JSON 字符串
        filename_mapping_json = json.dumps(filename_mapping, ensure_ascii=False)
        
        payload = {
            "inputs": {
                "numb": bianhao,
                "textPart3": text_part3,
                "evidence": evidence_files,
                "filename_mapping": filename_mapping_json
            },
            "response_mode": "blocking",
            "user": DIFY_USER_ID
        }
        
        logger.info(f"[AnalyzeEvidence] 开始调用Dify，证据文件数: {len(evidence_files)}")
        
        # 设置15分钟超时
        resp = requests.post(
            f"{DIFY_BASE_URL}/workflows/run",
            headers={"Authorization": f"Bearer {DIFY_API_KEY}", "Content-Type": "application/json"},
            json=payload,
            timeout=900
        )
        
        if resp.status_code != 200:
            error_text = resp.text[:500]
            logger.error(f"[AnalyzeEvidence] Dify错误: status={resp.status_code}, response={error_text}")
            return jsonify({'code': 500, 'message': f'Dify错误:{resp.status_code}'}), 500
        
        result = resp.json()
        logger.info(f"[AnalyzeEvidence] Dify返回成功")
        
        # 提取返回结果
        result_data = result.get('data', {})
        outputs = result_data.get('outputs', {})
        
        return jsonify({
            'code': 200,
            'success': True,
            'message': '证据分析完成',
            'data': {
                'outputs': outputs,
                'task_id': result.get('task_id', 'N/A'),
                'evidence_count': len(evidence_files)
            }
        })
        
    except requests.exceptions.Timeout:
        logger.error("[AnalyzeEvidence] Dify调用超时（超过15分钟）")
        return jsonify({'code': 504, 'message': '分析超时（超过15分钟），请稍后重试'}), 504
    except Exception as e:
        logger.error(f"[AnalyzeEvidence] 异常: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'code': 500, 'message': str(e)}), 500


@app.route('/api/award/status/<case_id>', methods=['GET'])
def award_status(case_id):
    """
    查询裁决书生成状态
    返回生成文件路径（如果有）
    """
    try:
        # 从案号提取编号
        import re
        bianhao = case_id
        
        # 先获取案件详情提取案号
        if login_manager.check_and_renew_login():
            headers = login_manager.get_auth_headers()
            detail_url = "http://10.96.10.78:8080/v1/api/admin/case/caseData"
            response = requests.post(
                detail_url,
                headers=headers,
                json={'id': case_id},
                timeout=30
            )
            if response.status_code == 200:
                case_data = response.json()
                data_content = case_data.get('data', {})
                if isinstance(data_content, dict):
                    case_detail = data_content
                elif isinstance(data_content, list) and len(data_content) > 0:
                    case_detail = data_content[0] if isinstance(data_content[0], dict) else {}
                else:
                    case_detail = {}
                
                case_no = case_detail.get('case_no', '')
                if case_no:
                    match = re.search(r'\[(\d{4})\](\d+)', case_no)
                    if match:
                        year, num = match.groups()
                        bianhao = f"{year}{num}"
        
        # 查询数据库
        conn = db_manager.get_connection()
        cursor = conn.cursor(dictionary=True)
        
        try:
            cursor.execute(
                "SELECT * FROM `裁决书要素保存` WHERE `案号` = %s",
                (bianhao,)
            )
            result = cursor.fetchone()
            
            if result:
                file_paths = result.get('生成文件路径', '')
                if file_paths:
                    # 文件已生成
                    return jsonify({
                        'code': 200,
                        'message': '裁决书已生成',
                        'data': {
                            'status': 'completed',
                            'file_paths': file_paths.split(','),
                            'generated_at': result.get('updated_at', '')
                        },
                        'timestamp': datetime.now().isoformat()
                    })
                else:
                    # 尚未生成
                    return jsonify({
                        'code': 200,
                        'message': '裁决书生成中，请稍后',
                        'data': {
                            'status': 'generating',
                            'file_paths': [],
                            'bianhao': bianhao
                        },
                        'timestamp': datetime.now().isoformat()
                    })
            else:
                return jsonify({
                    'code': 404,
                    'message': '未找到该案件的裁决书要素',
                    'data': None,
                    'timestamp': datetime.now().isoformat()
                })
        finally:
            cursor.close()
            conn.close()
            
    except Exception as e:
        logger.error(f"查询裁决书状态失败: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            'code': 500,
            'message': f'查询失败: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


@app.route('/api/award/download', methods=['GET'])
def download_award():
    """
    下载裁决书文件
    参数:
      - path: 文件路径（URL编码）
    """
    try:
        from urllib.parse import unquote
        
        # 获取文件路径参数
        file_path = request.args.get('path', '')
        if not file_path:
            return jsonify({
                'code': 400,
                'message': '缺少参数: path',
                'timestamp': datetime.now().isoformat()
            }), 400
        
        # URL解码
        file_path = unquote(file_path)
        
        # 安全检查：确保路径在项目目录下
        base_dir = os.path.dirname(os.path.abspath(__file__))
        full_path = os.path.join(base_dir, file_path)
        
        # 规范化路径并检查是否在允许范围内
        full_path = os.path.normpath(full_path)
        if not full_path.startswith(base_dir):
            logger.warning(f"尝试访问非法路径: {file_path}")
            return jsonify({
                'code': 403,
                'message': '访问被拒绝: 非法文件路径',
                'timestamp': datetime.now().isoformat()
            }), 403
        
        # 检查文件是否存在
        if not os.path.exists(full_path) or not os.path.isfile(full_path):
            logger.warning(f"文件不存在: {full_path}")
            return jsonify({
                'code': 404,
                'message': '文件不存在',
                'timestamp': datetime.now().isoformat()
            }), 404
        
        # 获取文件名
        file_name = os.path.basename(full_path)
        
        # 返回文件
        from flask import send_file
        return send_file(
            full_path,
            as_attachment=True,
            download_name=file_name,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"下载裁决书失败: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            'code': 500,
            'message': f'查询失败: {str(e)}',
            'data': None,
            'timestamp': datetime.now().isoformat()
        }), 500


# ============================================
# 启动服务器
# ============================================

def start_server():
    """启动服务器"""
    logger.info("=" * 60)
    logger.info("劳动仲裁信息查询综合服务平台启动")
    logger.info("=" * 60)
    logger.info(f"服务地址: http://{Config.FLASK_HOST}:{Config.FLASK_PORT}")
    logger.info(f"页面访问: http://<服务器IP>:{Config.FLASK_PORT}/")
    logger.info(f"案件查询: http://<服务器IP>:{Config.FLASK_PORT}/query")
    logger.info("=" * 60)
    logger.info("API端点列表:")
    logger.info("  [页面服务]")
    logger.info("  GET  /              - 劳动仲裁申请书在线填写（支持编辑: ?case_id=xxx）")
    logger.info("  GET  /query         - 案件查询页面")
    logger.info("  GET  /cases         - 案件管理列表")
    logger.info("  GET  /receive_query  - 收件查询页面")
    logger.info("  GET  /receive_detail - 收件详情页面")
    logger.info("  GET  /handle_query   - 立案查询页面")
    logger.info("  GET  /handle_detail  - 立案详情页面")
    logger.info("  [内部API服务]")
    logger.info("  GET  /api/status    - 服务状态")
    logger.info("  GET  /api/login/status    - 登录状态")
    logger.info("  POST /api/login           - 手动登录")
    logger.info("  POST /api/company/query   - 查询企业信息")
    logger.info("  POST /api/idcard/query    - 查询身份证信息")
    logger.info("  GET  /api/db/status       - 数据库状态")
    logger.info("  [案件管理API]")
    logger.info("  POST /api/cases/save      - 保存案件")
    logger.info("  GET  /api/cases/query     - 查询案件")
    logger.info("  GET  /api/cases/list      - 案件列表")
    logger.info("  DELETE /api/cases/<id>    - 删除案件")
    logger.info("  GET  /api/receive/query   - 收件查询（代理内部服务）")
    logger.info("  GET  /api/receive/detail  - 收件详情（代理内部服务）")
    logger.info("  GET  /api/handle/query    - 立案查询（代理内部服务）")
    logger.info("  GET  /api/handle/detail   - 立案详情（代理内部服务）")
    logger.info("  GET  /api/file/proxy      - 文件代理下载")
    logger.info("  GET  /api/doc_templates/tree     - 文档模板目录树")
    logger.info("  GET  /api/doc_templates/download - 下载文档模板")
    logger.info("  POST /api/doc_templates/generate - 生成文档")
    logger.info("  [裁决书制作API]")
    logger.info("  GET  /award/make                 - 裁决书制作页面")
    logger.info("  GET  /api/award/elements/<id>    - 获取裁决书要素")
    logger.info("  POST /api/award/elements/<id>    - 保存裁决书要素")
    logger.info("  POST /api/award/generate         - 生成裁决书Word")
    logger.info("=" * 60)
    
    app.run(
        host=Config.FLASK_HOST,
        port=Config.FLASK_PORT,
        debug=Config.FLASK_DEBUG,
        threaded=True
    )


if __name__ == '__main__':
    start_server()
